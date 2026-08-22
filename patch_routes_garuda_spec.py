from docx.shared import Cm, Pt
import re

def patch():
    with open("app/routes.py", "r") as f:
        code = f.read()

    # 1. Patch _add_garuda
    old_garuda = """def _add_garuda(doc):
    path = Path(__file__).resolve().parent / "static" / "img" / "garuda.png"
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(path), width=Cm(2.8))"""
        
    new_garuda = """def _add_garuda(doc):
    path = Path(__file__).resolve().parent / "static" / "img" / "garuda.png"
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str(path), width=Cm(2.8))
        # Convert to anchor
        from lxml import etree
        drawing = run._element.find('.//' + qn('w:drawing'))
        if drawing is not None:
            inline = drawing.find('.//' + qn('wp:inline'))
            if inline is not None:
                inline.tag = qn('wp:anchor')
                inline.set('behindDoc', '0')
                inline.set('locked', '0')
                inline.set('layoutInCell', '1')
                inline.set('allowOverlap', '1')
                inline.set('simplePos', '0')
                inline.set('relativeHeight', '251658240')
                
                # Add position elements
                simplePos = etree.Element(qn('wp:simplePos'), x="0", y="0")
                positionH = etree.Element(qn('wp:positionH'), relativeFrom="column")
                etree.SubElement(positionH, qn('wp:posOffset')).text = "0"
                positionV = etree.Element(qn('wp:positionV'), relativeFrom="paragraph")
                etree.SubElement(positionV, qn('wp:posOffset')).text = "635"
                wrapNone = etree.Element(qn('wp:wrapNone'))
                
                # Insert them in correct order
                inline.insert(0, simplePos)
                inline.insert(1, positionH)
                inline.insert(2, positionV)
                # find extent
                extent_idx = 3
                for i, child in enumerate(inline):
                    if child.tag == qn('wp:effectExtent'):
                        extent_idx = i + 1
                        break
                inline.insert(extent_idx, wrapNone)"""
                
    if old_garuda in code:
        code = code.replace(old_garuda, new_garuda)
        print("Patched _add_garuda")
    else:
        print("Could not find _add_garuda to patch")


    # 2. Patch _add_spec
    match = re.search(r"def _add_spec\(doc, purchase\):(.*?)def _add_acceptance_receipt\(doc, purchase\):", code, re.DOTALL)
    if not match:
        print("Could not find _add_spec!")
        return

    new_spec = """def _add_spec(doc, purchase):
    _add_garuda(doc)
    _paragraph(doc, "บันทึกข้อความ", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    _paragraph(doc, "")
    _paragraph(doc, "ส่วนราชการ\\tโรงพยาบาลสิงห์บุรี กลุ่มงานเภสัชกรรม โทร. ๐๓๖๕๒ ๒๕๐๘ ต่อ ๑๑๒๙")
    _paragraph(doc, f"ที่ สห ๐๐๓๓.๒๐๕.๑๒/........................\\t\\t\\tวันที่ {purchase.document_date.strftime('%d/%m/%Y')}")
    _paragraph(doc, f"เรื่อง\\tขออนุมัติแต่งตั้งผู้กำหนดรายละเอียดคุณลักษณะเฉพาะของพัสดุ จำนวน {len(purchase.lines)} รายการ")
    _paragraph(doc, "เรียน\\tผู้ว่าราชการจังหวัดสิงห์บุรี")
    _paragraph(doc, f"\\tด้วยกลุ่มงานเภสัชกรรม โรงพยาบาลสิงห์บุรี จะดำเนินการจัดซื้อเวชภัณฑ์มิใช่ยา จำนวน {len(purchase.lines)} รายการ  ดังนี้")
    _paragraph(doc, "")

    table = doc.add_table(rows=1, cols=5)
    _apply_table_grid(table)
    headers = ["ลำดับ", "รายการ", "จำนวน", "หน่วยละ", "เป็นเงิน"]
    table.autofit = False
    widths = [Cm(1.2), Cm(9.3), Cm(2.3), Cm(2.5), Cm(2.4)]
    for i in range(5):
        table.columns[i].width = widths[i]
    for i, value in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], value, True, WD_ALIGN_PARAGRAPH.CENTER, size=14)
        table.rows[0].cells[i].width = widths[i]
    for line in purchase.lines:
        cells = table.add_row().cells
        values = [line.line_no, line.description, f"{line.quantity:g} {line.unit.name}", f"{line.unit_price:,.2f}", f"{line.amount:,.2f}"]
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            _set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER, size=14)
    cells = table.add_row().cells
    _set_cell_text(cells[1], f"({baht_text(purchase.total_amount)})", True, WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _set_cell_text(cells[4], f"{purchase.total_amount:,.2f}", True, WD_ALIGN_PARAGRAPH.RIGHT, size=14)
    
    _paragraph(doc, "เพื่อให้ได้ร่างขอบเขตของงานหรือรายละเอียดคุณลักษณะเฉพาะของพัสดุดังกล่าว รวมทั้งกำหนดหลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอ เพื่อใช้ในการจัดซื้อ/จัดจ้าง ตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 21 วรรคหนึ่ง และมติคณะรัฐมนตรีและหลักเกณฑ์ที่เกี่ยวข้อง จึงขออนุมัติแต่งตั้งคณะกรรมการหรือผู้กำหนด หรือบุคคลใดบุคคลหนึ่ง ในการจัดทำร่างขอบเขตของงานหรือรายละเอียดคุณลักษณะเฉพาะของพัสดุที่จะซื้อ รวมทั้งกำหนดหลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอ ตามรายชื่อดังนี้", first_line=True)
    _paragraph(doc, "\\t1.\\t\\t\\tนางสาวนลินี เครือทิวา\\t\\t\\t\\tตำแหน่ง\\tเภสัชกรชำนาญการ")
    _paragraph(doc, "\\t\\t\\tจึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ")
    _paragraph(doc, "\\t\\t\\t\\t\\t\\tลงชื่อ ....................................................... เจ้าหน้าที่")
    _paragraph(doc, "\\t\\t\\t\\t\\t\\t(นางพิณนภา ศริพันธุ์)")
    _paragraph(doc, "\\t\\t\\t\\t\\t\\tเภสัชกรชำนาญการ")
    _paragraph(doc, "\\t\\t\\t\\t\\t\\tลงชื่อ ...................................................... หัวหน้าเจ้าหน้าที่")
    _paragraph(doc, "\\t\\t\\t\\t\\t\\t(นายชัชวาลย์ บุญญฤทธิ์)")
    _paragraph(doc, "\\t\\t\\t\\t\\t\\tเภสัชกรชำนาญการพิเศษ")
    _paragraph(doc, "\\t\\t\\t\\tอนุมัติ")
    _paragraph(doc, "\\t\\t(นายพิรุณ ปิตะหงษ์นันท์)")
    _paragraph(doc, "ผู้อำนวยการโรงพยาบาลสิงห์บุรี ปฏิบัติงานแทน")
    _paragraph(doc, "\\tผู้ว่าราชการจังหวัดสิงห์บุรี")
    
    doc.add_page_break()
    _paragraph(doc, "รายละเอียดคุณลักษณะเฉพาะ", WD_ALIGN_PARAGRAPH.CENTER, True, 20)
    _paragraph(doc, f"พัสดุจำนวน {len(purchase.lines)} รายการ")
    for line in purchase.lines:
        _paragraph(doc, f"{line.line_no}. {line.description}")
        _paragraph(doc, "คุณลักษณะเฉพาะตัวอย่าง สามารถแก้ไขได้")
        _paragraph(doc, f"จำนวน {line.quantity:g} {line.unit.name}")
    _paragraph(doc, "")
    _paragraph(doc, "\\tลงชื่อ ................................................ ผู้กำหนดรายละเอียด\\n\\t(................................................)", WD_ALIGN_PARAGRAPH.CENTER)
    
def _add_acceptance_receipt(doc, purchase):"""

    code = code[:match.start()] + new_spec + code[match.end()-38:]
    with open("app/routes.py", "w") as f:
        f.write(code)
    print("Patched _add_spec")

patch()
