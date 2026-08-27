import copy
import re
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path

from flask import Blueprint, flash, jsonify, redirect, render_template, request, send_file, url_for
from sqlalchemy import or_

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Emu

THAI_MONTHS_FULL = [
    "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
    "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
]

def format_thai_date_full(dt):
    if not dt: return ""
    return f"{dt.day} {THAI_MONTHS_FULL[dt.month - 1]} {dt.year + 543}"
from docx.text.run import Run as DocxRun
from openpyxl import Workbook as ExcelWorkbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


from .models import Company, DropdownOption, GovernmentProfile, Item, Purchase, PurchaseLine, Unit, db
from .utils import baht_text, safe_filename, to_decimal

main_bp = Blueprint("main", __name__)

# ตัวอักษร/สัญลักษณ์ทั่วไปใช้ฟอนต์ text, ตัวเลข (0-9) ใช้ฟอนต์ number
# (Word เลือกฟอนต์ตาม Unicode range ของตัวอักษรในแต่ละ run โดยอัตโนมัติ:
# w:ascii/w:hAnsi สำหรับเลข 0-9 และอักษรละติน, w:eastAsia/w:cs สำหรับอักษรไทย)
TEXT_FONT_NAME = "TH SarabunIT๙"
NUMBER_FONT_NAME = "TH Sarabun New"


@main_bp.app_template_filter("money")
def money(value):
    return f"{to_decimal(value):,.2f}"


@main_bp.app_context_processor
def inject_helpers():
    return {"baht_text": baht_text}


@main_bp.route("/")
def dashboard():
    return render_template(
        "dashboard.html",
        company_count=Company.query.filter_by(active=True).count(),
        item_count=Item.query.filter_by(active=True).count(),
        unit_count=Unit.query.filter_by(active=True).count(),
        purchase_count=Purchase.query.filter_by(active=True).count(),
        recent=Purchase.query.order_by(Purchase.id.desc()).limit(5).all(),
    )


# Companies
@main_bp.route("/companies")
def companies():
    q = request.args.get("q", "").strip()
    query = Company.query
    if q:
        query = query.filter(or_(Company.name.ilike(f"%{q}%"), Company.tax_id.ilike(f"%{q}%"), Company.phone.ilike(f"%{q}%")))
    return render_template("companies/list.html", companies=query.order_by(Company.name).all(), q=q)


@main_bp.route("/companies/new", methods=["GET", "POST"])
@main_bp.route("/companies/<int:company_id>/edit", methods=["GET", "POST"])
def company_form(company_id=None):
    company = Company.query.get_or_404(company_id) if company_id else Company()
    if request.method == "POST":
        company.name = request.form.get("name", "").strip()
        if not company.name:
            flash("กรุณากรอกชื่อบริษัท", "danger")
            return render_template("companies/form.html", company=company)
        for field in ["address", "phone", "tax_id", "bank_name", "bank_branch", "account_no", "account_name"]:
            setattr(company, field, request.form.get(field, "").strip())
        if not company_id:
            db.session.add(company)
        db.session.commit()
        flash("บันทึกข้อมูลบริษัทแล้ว ข้อมูลจะอยู่ถาวรในฐานข้อมูล", "success")
        return redirect(url_for("main.companies"))
    return render_template("companies/form.html", company=company)


@main_bp.post("/companies/<int:company_id>/toggle")
def company_toggle(company_id):
    company = Company.query.get_or_404(company_id)
    company.active = not company.active
    db.session.commit()
    flash("เปลี่ยนสถานะบริษัทแล้ว", "success")
    return redirect(url_for("main.companies"))


# Units
@main_bp.route("/units")
def units():
    q = request.args.get("q", "").strip()
    query = Unit.query
    if q:
        query = query.filter(Unit.name.ilike(f"%{q}%"))
    return render_template("units/list.html", units=query.order_by(Unit.name).all(), q=q)


@main_bp.route("/units/new", methods=["GET", "POST"])
@main_bp.route("/units/<int:unit_id>/edit", methods=["GET", "POST"])
def unit_form(unit_id=None):
    unit = Unit.query.get_or_404(unit_id) if unit_id else Unit()
    if request.method == "POST":
        unit.name = request.form.get("name", "").strip()
        if not unit.name:
            flash("กรุณากรอกชื่อหน่วย", "danger")
            return render_template("units/form.html", unit=unit)
        duplicate = Unit.query.filter(Unit.name == unit.name, Unit.id != (unit.id or 0)).first()
        if duplicate:
            flash("มีหน่วยนี้อยู่แล้ว", "danger")
            return render_template("units/form.html", unit=unit)
        if not unit_id:
            db.session.add(unit)
        db.session.commit()
        flash("บันทึกข้อมูลหน่วยแล้ว ข้อมูลจะอยู่ถาวรในฐานข้อมูล", "success")
        return redirect(url_for("main.units"))
    return render_template("units/form.html", unit=unit)


@main_bp.post("/units/<int:unit_id>/toggle")
def unit_toggle(unit_id):
    unit = Unit.query.get_or_404(unit_id)
    unit.active = not unit.active
    db.session.commit()
    flash("เปลี่ยนสถานะหน่วยแล้ว", "success")
    return redirect(url_for("main.units"))


# Items
@main_bp.route("/items")
def items():
    q = request.args.get("q", "").strip()
    query = Item.query
    if q:
        query = query.filter(or_(Item.name.ilike(f"%{q}%"), Item.code.ilike(f"%{q}%")))
    return render_template("items/list.html", items=query.order_by(Item.name).all(), q=q)


@main_bp.route("/items/new", methods=["GET", "POST"])
@main_bp.route("/items/<int:item_id>/edit", methods=["GET", "POST"])
def item_form(item_id=None):
    item = Item.query.get_or_404(item_id) if item_id else Item()
    units = Unit.query.filter_by(active=True).order_by(Unit.name).all()
    if request.method == "POST":
        item.code = request.form.get("code", "").strip()
        item.name = request.form.get("name", "").strip()
        item.specification = request.form.get("specification", "").strip()
        item.default_price = to_decimal(request.form.get("default_price"))
        item.unit_id = request.form.get("unit_id", type=int)
        if not item.name or not item.unit_id:
            flash("กรุณากรอกชื่อเวชภัณฑ์และหน่วย", "danger")
            return render_template("items/form.html", item=item, units=units)
        if not item_id:
            db.session.add(item)
        db.session.commit()
        flash("บันทึกข้อมูลเวชภัณฑ์แล้ว ข้อมูลจะอยู่ถาวรในฐานข้อมูล", "success")
        return redirect(url_for("main.items"))
    return render_template("items/form.html", item=item, units=units)


@main_bp.post("/items/<int:item_id>/toggle")
def item_toggle(item_id):
    item = Item.query.get_or_404(item_id)
    item.active = not item.active
    db.session.commit()
    flash("เปลี่ยนสถานะเวชภัณฑ์แล้ว", "success")
    return redirect(url_for("main.items"))


@main_bp.get("/api/items/<int:item_id>")
def item_api(item_id):
    item = Item.query.get_or_404(item_id)
    return jsonify({
        "id": item.id,
        "name": item.name,
        "price": f"{item.default_price:.2f}",
        "unit_id": item.unit_id,
        "unit_name": item.unit.name,
        "specification": item.specification or "",
    })


# Master dropdown data
@main_bp.route("/masters")
def masters():
    profiles = GovernmentProfile.query.order_by(GovernmentProfile.id).all()
    options = DropdownOption.query.order_by(DropdownOption.category, DropdownOption.id).all()
    return render_template("masters/list.html", profiles=profiles, options=options)


@main_bp.route("/masters/profile/new", methods=["GET", "POST"])
@main_bp.route("/masters/profile/<int:profile_id>/edit", methods=["GET", "POST"])
def profile_form(profile_id=None):
    profile = GovernmentProfile.query.get_or_404(profile_id) if profile_id else GovernmentProfile()
    if request.method == "POST":
        for field in ["name","department","letter_prefix","recipient","officer_name","officer_position","chief_name","chief_position","approver_name","approver_position","inspector1_name","inspector1_position","inspector2_name","inspector2_position","inspector3_name","inspector3_position","specifier_name","specifier_position","receipt1_name","receipt2_name","receipt3_name"]:
            setattr(profile, field, request.form.get(field, "").strip())
        if not profile_id:
            db.session.add(profile)
        db.session.commit()
        flash("บันทึกข้อมูลราชการแล้ว", "success")
        return redirect(url_for("main.masters"))
    return render_template("masters/profile_form.html", profile=profile)


@main_bp.route("/masters/option/new", methods=["GET", "POST"])
@main_bp.route("/masters/option/<int:option_id>/edit", methods=["GET", "POST"])
def option_form(option_id=None):
    option = DropdownOption.query.get_or_404(option_id) if option_id else DropdownOption()
    if request.method == "POST":
        option.category = request.form.get("category", "").strip()
        option.value = request.form.get("value", "").strip()
        if not option.category or not option.value:
            flash("กรุณากรอกหมวดและค่า", "danger")
        else:
            if not option_id:
                db.session.add(option)
            db.session.commit()
            flash("บันทึกตัวเลือก Dropdown แล้ว", "success")
            return redirect(url_for("main.masters"))
    return render_template("masters/option_form.html", option=option)


@main_bp.post("/masters/option/<int:option_id>/toggle")
def option_toggle(option_id):
    option = DropdownOption.query.get_or_404(option_id)
    option.active = not option.active
    db.session.commit()
    return redirect(url_for("main.masters"))


# Purchases
@main_bp.route("/purchases")
def purchases():
    q = request.args.get("q", "").strip()
    query = Purchase.query.join(Company)
    if q:
        query = query.filter(or_(Purchase.po_number.ilike(f"%{q}%"), Purchase.project_number.ilike(f"%{q}%"), Company.name.ilike(f"%{q}%")))
    return render_template("purchases/list.html", purchases=query.order_by(Purchase.document_date.desc(), Purchase.id.desc()).all(), q=q)


def purchase_context(purchase):
    categories = {}
    for option in DropdownOption.query.filter_by(active=True).order_by(DropdownOption.category, DropdownOption.id).all():
        categories.setdefault(option.category, []).append(option)
    return {
        "purchase": purchase,
        "companies": Company.query.filter_by(active=True).order_by(Company.name).all(),
        "items": Item.query.filter_by(active=True).order_by(Item.name).all(),
        "units": Unit.query.filter_by(active=True).order_by(Unit.name).all(),
        "government_profiles": GovernmentProfile.query.filter_by(active=True).order_by(GovernmentProfile.name).all(),
        "dropdowns": categories,
        "today": date.today().isoformat(),
    }


@main_bp.route("/purchases/new", methods=["GET", "POST"])
@main_bp.route("/purchases/<int:purchase_id>/edit", methods=["GET", "POST"])
def purchase_form(purchase_id=None):
    purchase = Purchase.query.get_or_404(purchase_id) if purchase_id else Purchase(document_date=date.today())
    if request.method == "POST":
        po_number = request.form.get("po_number", "").strip()
        duplicate = Purchase.query.filter(Purchase.po_number == po_number, Purchase.id != (purchase.id or 0)).first()
        if not po_number or duplicate:
            flash("กรุณากรอกเลขที่ใบสั่งซื้อที่ไม่ซ้ำ", "danger")
            return render_template("purchases/form.html", **purchase_context(purchase))

        purchase.po_number = po_number
        try:
            purchase.document_date = datetime.strptime(request.form.get("document_date", ""), "%Y-%m-%d").date()
        except ValueError:
            purchase.document_date = date.today()
        purchase.company_id = request.form.get("company_id", type=int)
        purchase.government_profile_id = request.form.get("government_profile_id", type=int)
        purchase.procurement_type = request.form.get("procurement_type", "").strip() or "เวชภัณฑ์มิใช่ยา"
        purchase.necessity_reason = request.form.get("necessity_reason", "").strip() or "ใช้ในการรักษาผู้ป่วย"
        purchase.project_number = request.form.get("project_number", "").strip()
        purchase.contract_control_number = request.form.get("contract_control_number", "").strip()
        purchase.delivery_days = request.form.get("delivery_days", type=int) or 30
        purchase.delivery_place = request.form.get("delivery_place", "").strip()
        purchase.budget_source = request.form.get("budget_source", "").strip()
        purchase.budget_allocated = to_decimal(request.form.get("budget_allocated"))
        purchase.budget_previously_used = to_decimal(request.form.get("budget_previously_used"))
        purchase.note = request.form.get("note", "").strip()

        lines = []
        for index in range(1, 7):
            item_id = request.form.get(f"item_id_{index}", type=int)
            if not item_id:
                continue
            item = db.session.get(Item, item_id)
            if item is None:
                continue
            quantity = to_decimal(request.form.get(f"quantity_{index}"))
            unit_price = to_decimal(request.form.get(f"unit_price_{index}"))
            unit_id = request.form.get(f"unit_id_{index}", type=int) or item.unit_id
            if quantity <= 0:
                continue
            lines.append(PurchaseLine(
                line_no=len(lines) + 1,
                item_id=item.id,
                description=request.form.get(f"description_{index}", "").strip() or item.name,
                specification=request.form.get(f"specification_{index}", "").strip() or item.specification,
                quantity=quantity,
                unit_id=unit_id,
                unit_price=unit_price,
                amount=(quantity * unit_price).quantize(Decimal("0.01")),
            ))

        if not purchase.company_id or not lines:
            flash("กรุณาเลือกบริษัทและเพิ่มอย่างน้อย 1 รายการ", "danger")
            return render_template("purchases/form.html", **purchase_context(purchase))

        if not purchase_id:
            db.session.add(purchase)
        purchase.lines.clear()
        purchase.lines.extend(lines)
        db.session.commit()
        flash("บันทึกใบสั่งซื้อแล้ว ข้อมูลถูกเก็บในฐานข้อมูลถาวร สามารถส่งออก Word ได้ด้านล่าง", "success")
        return redirect(url_for("main.purchase_detail", purchase_id=purchase.id))

    return render_template("purchases/form.html", **purchase_context(purchase))


@main_bp.route("/purchases/<int:purchase_id>")
def purchase_detail(purchase_id):
    return render_template("purchases/detail.html", purchase=Purchase.query.get_or_404(purchase_id))


@main_bp.post("/purchases/<int:purchase_id>/toggle")
def purchase_toggle(purchase_id):
    purchase = Purchase.query.get_or_404(purchase_id)
    purchase.active = not purchase.active
    db.session.commit()
    flash("เปลี่ยนสถานะเอกสารแล้ว", "success")
    return redirect(url_for("main.purchases"))


def _vat_values(purchase):
    total = to_decimal(purchase.total_amount)
    subtotal = (total / Decimal("1.07")).quantize(Decimal("0.01")) if total else Decimal("0.00")
    vat = (total - subtotal).quantize(Decimal("0.01"))
    return subtotal, vat


@main_bp.route("/purchases/<int:purchase_id>/print/<form_type>")
def print_form(purchase_id, form_type):
    purchase = Purchase.query.get_or_404(purchase_id)
    if form_type == "po":
        template = "print/purchase_single.html" if len(purchase.lines) == 1 else "print/purchase_multiple.html"
    elif form_type == "spec":
        template = "print/spec.html"
    elif form_type == "all":
        template = "print/all.html"
    else:
        return "ไม่พบแบบฟอร์ม", 404
    subtotal, vat = _vat_values(purchase)
    return render_template(template, purchase=purchase, subtotal_before_vat=subtotal, vat_amount=vat, embedded=False)


# ---------------- Word export ----------------
def _set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=16):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text or ""))
    run.bold = bold
    _set_run_font(run, size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return p


_DIGIT_SPAN_RE = re.compile(r"[0-9]+(?:[,.][0-9]+)*")


def _split_digit_segments(text):
    """แบ่งข้อความเป็นช่วง ๆ ตาม kind: "digit" (เลข 0-9 ล้วน), "break" (ตัวแทน \\n/\\r
    ที่ python-docx ใช้แทน <w:br/> ตอนอ่าน run.text), "tab" (ตัวแทน \\t ของ <w:tab/>),
    หรือ "text" (ตัวอักษร/สัญลักษณ์อื่นทั้งหมด รวมตัวอักษรอังกฤษ A-Z).

    ต้องแยก \\n/\\t ออกมาต่างหาก ไม่ปนไปกับ "text" เพราะไม่งั้นตอนเขียนกลับจะกลาย
    เป็นตัวอักษร \\n ดิบฝังใน <w:t> ซึ่ง Word ไม่ตีความเป็นการขึ้นบรรทัดใหม่ (layout พัง)."""
    text = text or ""
    segments = []
    buf = ""

    def flush_text():
        nonlocal buf
        if not buf:
            return
        pos = 0
        for m in _DIGIT_SPAN_RE.finditer(buf):
            if m.start() > pos:
                segments.append((buf[pos:m.start()], "text"))
            segments.append((m.group(), "digit"))
            pos = m.end()
        if pos < len(buf):
            segments.append((buf[pos:], "text"))
        buf = ""

    for ch in text:
        if ch in ("\n", "\r"):
            flush_text()
            segments.append((ch, "break"))
        elif ch == "\t":
            flush_text()
            segments.append((ch, "tab"))
        else:
            buf += ch
    flush_text()
    return segments or [("", "text")]


def _set_run_lang(rpr):
    """Mark the run as Thai complex-script text so Word actually reads the
    w:cs font (otherwise the Font Name box falls back to showing w:ascii)."""
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:eastAsia"), "th-TH")
    lang.set(qn("w:bidi"), "th-TH")


def _apply_font_to_rpr(rpr, is_digit):
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    ascii_font = NUMBER_FONT_NAME if is_digit else TEXT_FONT_NAME
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)


_RUN_CONTENT_TAGS = (qn("w:t"), qn("w:br"), qn("w:tab"), qn("w:cr"))


_digit_font_split_disabled = [False]


def _digit_font_split_active():
    return not _digit_font_split_disabled[-1]


class _DisableDigitFontSplit:
    """ปิดการแยกฟอนต์เลข/ข้อความชั่วคราว ใช้ครอบ _build_exact_procurement_template
    เพราะ template ราชการต้นฉบับ (purchase_master.docx) จัดระยะตัวอักษรมาละเอียด
    ทั้งเอกสารด้วยฟอนต์เดียว การแตก run เพื่อแยกฟอนต์เลขในนั้นทำให้ตัวอักษรดูยืด/
    เพี้ยนเมื่อเปิดใน Word (ต่างจากฟอร์มอื่นที่สร้างจาก python-docx เปล่า ๆ ซึ่งไม่มี
    การจัดระยะละเอียดแบบนี้ จึงแยกฟอนต์ได้ตามปกติ)."""

    def __enter__(self):
        _digit_font_split_disabled.append(True)
        return self

    def __exit__(self, *exc_info):
        _digit_font_split_disabled.pop()


def _write_split_text(run_element, text):
    """เขียนข้อความลง run_element โดยแยกเป็นหลาย run เมื่อมีทั้งเลขและตัวอักษรปนกัน
    เพื่อให้เลข 0-9 ใช้ NUMBER_FONT_NAME ส่วนตัวอักษร/สัญลักษณ์อื่นใช้ TEXT_FONT_NAME
    และคง \\n/\\t (จาก python-docx run.text) เป็น <w:br/>/<w:tab/> เหมือนเดิม แทนที่จะ
    ฝังเป็นตัวอักษรดิบใน <w:t> ซึ่งจะทำให้บรรทัดพังตอนเปิดใน Word."""
    segments = _split_digit_segments(text)

    def _clear_content(element):
        for child in list(element):
            if child.tag in _RUN_CONTENT_TAGS:
                element.remove(child)

    def _append_segment(element, value, kind):
        if kind == "break":
            element.append(OxmlElement("w:br"))
        elif kind == "tab":
            element.append(OxmlElement("w:tab"))
        else:
            t_el = OxmlElement("w:t")
            t_el.text = value
            if value != value.strip():
                t_el.set(qn("xml:space"), "preserve")
            element.append(t_el)

    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_element.insert(0, rpr)

    if not _digit_font_split_active():
        # เปลี่ยนทุก segment ที่เป็น "digit" ให้ใช้ฟอนต์ข้อความแทน (ไม่แยกเลข) เพื่อ
        # รักษา layout ที่จัดมาละเอียดของ template ไว้ (ดู _DisableDigitFontSplit
        # ด้านบน) แต่ "break"/"tab" ต้องคงไว้เหมือนเดิม ไม่งั้น \n/\t จากต้นทาง (เช่น
        # python-docx run.text ที่แปลง <w:br/> เป็น "\n") จะถูกฝังเป็นตัวอักษรดิบ
        # ใน <w:t> แทนที่จะเป็น <w:br/>/<w:tab/> ทำให้บรรทัดพังเหมือนบั๊กที่เจอมาก่อน
        segments = [
            (seg_text, "text" if kind == "digit" else kind)
            for seg_text, kind in segments
        ]
    else:
        # run ที่มี character scaling (w:w) มักมาจาก template ราชการที่ปรับระยะ/
        # ความกว้างตัวอักษรมาละเอียดให้พอดีบรรทัดอยู่แล้ว การแตก run เป็นหลายชิ้นจะ
        # ทำลาย kerning ระหว่างรอยต่อ ทำให้ตัวอักษรดูยืด/เพี้ยนเมื่อเปิดใน Word จึงไม่
        # แตก run ในกรณีนี้ ยอมให้เลขใช้ฟอนต์ข้อความปนไปแทน เพื่อรักษาระยะห่างเดิมไว้
        w_el = rpr.find(qn("w:w"))
        has_char_scale = w_el is not None and w_el.get(qn("w:val")) not in (None, "100")
        if has_char_scale and len(segments) > 1:
            segments = [(text, "text")]

    _clear_content(run_element)
    first_text, first_kind = segments[0]
    _apply_font_to_rpr(rpr, first_kind == "digit")
    _append_segment(run_element, first_text, first_kind)

    insert_after = run_element
    for segment_text, kind in segments[1:]:
        new_el = copy.deepcopy(run_element)
        _clear_content(new_el)
        _apply_font_to_rpr(new_el.find(qn("w:rPr")), kind == "digit")
        _append_segment(new_el, segment_text, kind)
        insert_after.addnext(new_el)
        insert_after = new_el


def _set_run_font(run, size=16, bold=None):
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    _write_split_text(run._element, run.text)


def _paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=16, first_line=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(str(text))
    _set_run_font(run, size, bold)
    return p


def _configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = TEXT_FONT_NAME
    normal.font.size = Pt(16)
    normal_rfonts = normal._element.rPr.rFonts
    normal_rfonts.set(qn("w:ascii"), TEXT_FONT_NAME)
    normal_rfonts.set(qn("w:hAnsi"), TEXT_FONT_NAME)
    normal_rfonts.set(qn("w:eastAsia"), TEXT_FONT_NAME)
    normal_rfonts.set(qn("w:cs"), TEXT_FONT_NAME)
    _set_run_lang(normal._element.rPr)


def _add_garuda(doc, width=None, height=None):
    path = Path(__file__).resolve().parent / "static" / "img" / "garuda.png"
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Keep the heading immediately below the Garuda.  The customer
        # specifically requested that memo headings not drift down the page.
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        # The document uses a document grid (w:docGrid linePitch) for
        # consistent printed line spacing, so Word still snaps this
        # picture-only line to the grid's line pitch and shows a dead gap
        # underneath it even with spacing at zero above. Disabling grid
        # snap here matches what the master template's own Garuda
        # paragraphs already do (see _xml_disable_snap_to_grid).
        _xml_disable_snap_to_grid(p._element)
        run = p.add_run()
        if width is not None and height is not None:
            run.add_picture(str(path), width=width, height=height)
        else:
            run.add_picture(str(path), width=Emu(635726), height=Emu(696026))

def _apply_table_grid(table):
    """Apply Table Grid when the template contains that built-in style, fallback to oxml borders."""
    try:
        table.style = "Table Grid"
    except KeyError:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        table._tbl.tblPr.append(tblBorders)


def _format_po_meta_cell(cell):
    """Keep seller and PO/header columns aligned from the same top baseline.

    The two cells contain different amounts of text.  Centering them vertically
    makes the shorter PO column start lower than the seller column, which is the
    layout problem marked by the customer.
    """
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0


def _add_purchase_order(doc, purchase):
    _add_garuda(doc, width=Cm(1.27), height=Cm(1.4))
    _paragraph(doc, "ใบสั่งซื้อ", WD_ALIGN_PARAGRAPH.CENTER, True, 20)

    meta = doc.add_table(rows=1, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta.columns[0].width = Cm(9.5)
    meta.columns[1].width = Cm(8.0)
    for row in meta.rows:
        row.cells[0].width = Cm(9.5)
        row.cells[1].width = Cm(8.0)
    left = (
        f"ผู้ขาย {purchase.company.name}\n"
        f"ที่อยู่ {purchase.company.address or '-'}\n"
        f"โทรศัพท์ {purchase.company.phone or '-'}\n"
        f"เลขประจำตัวผู้เสียภาษี {purchase.company.tax_id or '-'}\n"
        f"เลขที่บัญชีเงินฝากธนาคาร {purchase.company.account_no or '-'}\n"
        f"ชื่อบัญชี {purchase.company.account_name or '-'}\n"
        f"ธนาคาร {purchase.company.bank_name or '-'}"
        + (f" สาขา {purchase.company.bank_branch}" if purchase.company.bank_branch else "")
    )
    right = (
        f"เลขที่ {purchase.po_number}\n"
        f"วันที่ {format_thai_date_full(purchase.document_date)}\n\n"
        "ส่วนราชการ โรงพยาบาลสิงห์บุรี\n"
        "ที่อยู่ ๙๑๗/๓ ตำบลบางพุทรา อำเภอเมืองสิงห์บุรี จังหวัดสิงห์บุรี ๑๖๐๐๐\n"
        "โทรศัพท์ ๐๓๖-๕๒๒๕๐๗"
    )
    _set_cell_text(meta.cell(0, 0), left)
    _set_cell_text(meta.cell(0, 1), right)
    _format_po_meta_cell(meta.cell(0, 0))
    _format_po_meta_cell(meta.cell(0, 1))

    _paragraph(doc, f"ตามที่ {purchase.company.name} ได้เสนอราคาไว้ต่อโรงพยาบาลสิงห์บุรี ซึ่งได้รับราคาและตกลงซื้อตามรายการดังต่อไปนี้", first_line=True)

    table = doc.add_table(rows=1, cols=6)
    _apply_table_grid(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ลำดับ", "รายการ", "จำนวน", "หน่วย", "ราคาต่อหน่วย\n(บาท)", "จำนวนเงิน\n(บาท)"]
    table.autofit = False
    widths = [Cm(1.3), Cm(4.5), Cm(2.9), Cm(2.9), Cm(3.1), Cm(2.9)]
    for i in range(6):
        table.columns[i].width = widths[i]
    for i, value in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], value, True, WD_ALIGN_PARAGRAPH.CENTER)
        table.rows[0].cells[i].width = widths[i]
    for line in purchase.lines:
        cells = table.add_row().cells
        values = [line.line_no, line.description, f"{line.quantity:g}", line.unit.name, f"{line.unit_price:,.2f}", f"{line.amount:,.2f}"]
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else (WD_ALIGN_PARAGRAPH.RIGHT if i >= 4 else WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(cells[i], value, align=align)

    subtotal, vat = _vat_values(purchase)
    for label, value in [("รวมเป็นเงิน", subtotal), ("ภาษีมูลค่าเพิ่ม", vat), ("รวมเป็นเงินทั้งสิ้น", purchase.total_amount)]:
        cells = table.add_row().cells
        merged = cells[0].merge(cells[3])
        _set_cell_text(merged, f"({baht_text(purchase.total_amount)})" if label == "รวมเป็นเงิน" else "", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[4], label, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[5], f"{value:,.2f}", bold=(label == "รวมเป็นเงินทั้งสิ้น"), align=WD_ALIGN_PARAGRAPH.RIGHT)

    _paragraph(doc, "การซื้ออยู่ภายใต้เงื่อนไขต่อไปนี้", bold=True)
    conditions = [
        f"๑. กำหนดส่งมอบภายใน {purchase.delivery_days} วัน นับถัดจากวันที่ผู้ขายได้รับใบสั่งซื้อ",
        "๒. ครบกำหนดส่งมอบวันที่ ................................................................",
        f"๓. สถานที่ส่งมอบ {purchase.delivery_place or 'โรงพยาบาลสิงห์บุรี ๙๑๗/๓'}",
        "๔. ระยะเวลารับประกัน -",
        "๕. สงวนสิทธิ์ค่าปรับกรณีส่งมอบเกินกำหนด โดยคิดค่าปรับเป็นรายวันในอัตราร้อยละ ๐.๒๐ ของราคาสิ่งของที่ยังไม่ได้รับมอบ",
        "๖. ส่วนราชการสงวนสิทธิ์ที่จะไม่รับมอบ หากสินค้ามีลักษณะไม่ตรงตามรายการที่ระบุไว้ในใบสั่งซื้อ ผู้ขายจะต้องดำเนินการเปลี่ยนใหม่ให้ถูกต้องทุกประการ",
        "PAGE_BREAK",
        "๗. หน่วยงานของรัฐสามารถนำผลการปฏิบัติงานตามสัญญาหรือข้อตกลงมาประเมินผลการปฏิบัติงานของผู้ประกอบการ",
    ]
    for text in conditions:
        if text == "PAGE_BREAK":
            doc.add_page_break()
        else:
            _paragraph(doc, text)

    _paragraph(doc, "หมายเหตุ", bold=True)
    _paragraph(doc, "	๑. การติดอากรแสตมป์ให้เป็นไปตามประมวลกฎหมายรัษฎากร หากต้องการให้ใบสั่งซื้อมีผลตามกฎหมาย")
    _paragraph(doc, f"	๒. ใบสั่งซื้อนี้อ้างอิงตามเลขที่โครงการ {purchase.project_number or '........................'} ซื้อพัสดุจำนวน {len(purchase.lines)} รายการ เป็นเงิน {purchase.total_amount:,.2f} บาท ({baht_text(purchase.total_amount)}) โดยวิธีเฉพาะเจาะจง")
    _paragraph(doc, "")
    _paragraph(doc, "")
    _paragraph(doc, "\t\t\t\tลงชื่อ ................................................ ผู้สั่งซื้อ\n\t\t\t\t(................................................)\n\t\t\t\tหัวหน้าเจ้าหน้าที่\n\t\t\t\tวันที่ ................................................", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "")
    _paragraph(doc, "")
    _paragraph(doc, "\t\t\t\t\tลงชื่อ ................................................ ผู้รับใบสั่งซื้อ\n\t\t\t\t(................................................)\n\t\t\t\tวันที่ ................................................", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "")
    _paragraph(doc, f"เลขที่โครงการ {purchase.project_number or '........................'}")
    _paragraph(doc, f"เลขคุมสัญญา {purchase.contract_control_number or '........................'}")

def _add_spec(doc, purchase):
    _add_garuda(doc)
    _paragraph(doc, "บันทึกข้อความ", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    _paragraph(doc, "")
    _paragraph(doc, "ส่วนราชการ\tโรงพยาบาลสิงห์บุรี กลุ่มงานเภสัชกรรม โทร. ๐๓๖๕๒ ๒๕๐๘ ต่อ ๑๑๒๙")
    _paragraph(doc, f"ที่ สห ๐๐๓๓.๒๐๕.๑๒/........................\t\t\tวันที่ {format_thai_date_full(purchase.document_date)}")
    _paragraph(doc, f"เรื่อง\tขออนุมัติแต่งตั้งผู้กำหนดรายละเอียดคุณลักษณะเฉพาะของพัสดุ จำนวน {len(purchase.lines)} รายการ")
    _paragraph(doc, "เรียน\tผู้ว่าราชการจังหวัดสิงห์บุรี")
    _paragraph(doc, f"\tด้วยกลุ่มงานเภสัชกรรม โรงพยาบาลสิงห์บุรี จะดำเนินการจัดซื้อเวชภัณฑ์มิใช่ยา จำนวน {len(purchase.lines)} รายการ  ดังนี้")
    _paragraph(doc, "")

    table = doc.add_table(rows=1, cols=5)
    _apply_table_grid(table)
    headers = ["ลำดับ", "รายการ", "จำนวน", "หน่วยละ", "เป็นเงิน"]
    table.autofit = False
    widths = [Cm(1.2), Cm(9.3), Cm(2.3), Cm(2.5), Cm(2.4)]
    for i in range(5):
        table.columns[i].width = widths[i]
    for i, value in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], value, True, WD_ALIGN_PARAGRAPH.CENTER, size=14)
        table.rows[0].cells[i].width = widths[i]
    for line in purchase.lines:
        cells = table.add_row().cells
        values = [line.line_no, line.description, f"{line.quantity:g} {line.unit.name}", f"{line.unit_price:,.2f}", f"{line.amount:,.2f}"]
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            _set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER, size=14)
    cells = table.add_row().cells
    _set_cell_text(cells[1], f"({baht_text(purchase.total_amount)})", True, WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _set_cell_text(cells[4], f"{purchase.total_amount:,.2f}", True, WD_ALIGN_PARAGRAPH.RIGHT, size=14)
    
    _paragraph(doc, "เพื่อให้ได้ร่างขอบเขตของงานหรือรายละเอียดคุณลักษณะเฉพาะของพัสดุดังกล่าว รวมทั้งกำหนดหลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอ เพื่อใช้ในการจัดซื้อ/จัดจ้าง ตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 21 วรรคหนึ่ง และมติคณะรัฐมนตรีและหลักเกณฑ์ที่เกี่ยวข้อง จึงขออนุมัติแต่งตั้งคณะกรรมการหรือผู้กำหนด หรือบุคคลใดบุคคลหนึ่ง ในการจัดทำร่างขอบเขตของงานหรือรายละเอียดคุณลักษณะเฉพาะของพัสดุที่จะซื้อ รวมทั้งกำหนดหลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอ ตามรายชื่อดังนี้", first_line=True)
    _paragraph(doc, "\t1.\t\t\tนางสาวนลินี เครือทิวา\t\t\t\tตำแหน่ง\tเภสัชกรชำนาญการ")
    _paragraph(doc, "")
    _paragraph(doc, "\t\tจึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ")
    _paragraph(doc, "")
    _paragraph(doc, "\t\tลงชื่อ ....................................................... เจ้าหน้าที่", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "\t\t(นางพิณนภา ศริพันธุ์)", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "\t\tเภสัชกรชำนาญการ", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "")
    _paragraph(doc, "")
    _paragraph(doc, "\t\tลงชื่อ ...................................................... หัวหน้าเจ้าหน้าที่", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "\t\t(นายชัชวาลย์ บุญญฤทธิ์)", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "\t\tเภสัชกรชำนาญการพิเศษ", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "                       อนุมัติ")
    _paragraph(doc, "")
    _paragraph(doc, "")
    _paragraph(doc, "              (นายพิรุณ ปิตะหงษ์นันท์)")
    _paragraph(doc, "ผู้อำนวยการโรงพยาบาลสิงห์บุรี ปฏิบัติงานแทน")
    _paragraph(doc, "            ผู้ว่าราชการจังหวัดสิงห์บุรี")
    
    doc.add_page_break()
    _paragraph(doc, "รายละเอียดคุณลักษณะเฉพาะ", WD_ALIGN_PARAGRAPH.CENTER, True, 20)
    _paragraph(doc, f"พัสดุจำนวน {len(purchase.lines)} รายการ")
    for line in purchase.lines:
        _paragraph(doc, f"{line.line_no}. {line.description}")
        _paragraph(doc, "คุณลักษณะเฉพาะตัวอย่าง สามารถแก้ไขได้")
        _paragraph(doc, f"จำนวน {line.quantity:g} {line.unit.name}")
    _paragraph(doc, "")
    _paragraph(doc, "\t\t\t\t\tลงชื่อ ................................................ ผู้กำหนดรายละเอียด\n\t\t\t(................................................)", WD_ALIGN_PARAGRAPH.CENTER)
    
def _add_acceptance_receipt(doc, purchase):
    _paragraph(doc, "ใบตรวจรับการจัดซื้อ/จัดจ้าง", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    _paragraph(doc, f"\t\tวันที่ {format_thai_date_full(purchase.document_date)}", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(
        doc,
        f"ตามใบสั่งซื้อ เลขที่ {purchase.po_number} ลงวันที่ {format_thai_date_full(purchase.document_date)} "
        f"โรงพยาบาลสิงห์บุรีได้ตกลงซื้อกับ {purchase.company.name} สำหรับโครงการซื้อพัสดุ "
        f"จำนวน {len(purchase.lines)} รายการ โดยวิธีเฉพาะเจาะจง เป็นจำนวนเงินทั้งสิ้น "
        f"{purchase.total_amount:,.2f} บาท ({baht_text(purchase.total_amount)})",
        first_line=True,
    )
    _paragraph(doc, "")
    _paragraph(doc, "คณะกรรมการตรวจรับพัสดุ ได้ตรวจรับงานแล้ว ผลปรากฏดังนี้", first_line=True)
    _paragraph(doc, "")
    _paragraph(doc, "\t1. ผลการตรวจรับ")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    for part in ["\t", "☐", " ถูกต้อง      ", "☐", " ครบถ้วนตามสัญญา      ", "☐", " ไม่ครบถ้วนตามสัญญา"]:
        r = p.add_run(part)
        if part == "☐":
            r.font.size = Pt(14)
    _paragraph(doc, "")
    _paragraph(doc, "\t2. ค่าปรับ")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    for part in ["\t", "☐", " มีค่าปรับ      ", "☐", " ไม่มีค่าปรับ"]:
        r = p.add_run(part)
        if part == "☐":
            r.font.size = Pt(14)
    _paragraph(doc, "")
    _paragraph(doc, "\t3. การเบิกจ่ายเงิน")
    if len(purchase.lines) == 1:
        line = purchase.lines[0]
        _paragraph(doc, f"\tเบิกจ่ายเงิน เป็นจำนวนเงินทั้งสิ้น {line.amount:,.2f} บาท")
    else:
        for line in purchase.lines:
            _paragraph(doc, f"\t- รายการที่ {line.line_no} {line.description}")
            _paragraph(doc, f"\tเบิกจ่ายเงิน งวดที่ 1 เป็นจำนวนเงินทั้งสิ้น {line.amount:,.2f} บาท")

    _paragraph(doc, "")
    sig = doc.add_table(rows=3, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    names = [
        "(นางสาวธัญรดา ใจเสน)",
        "(นางสาวจุฑามาศ อาคมสรรเสริญ)",
        "(นางสาวนริศรา ม่วงงาม)",
    ]
    for i, name in enumerate(names):
        _set_cell_text(sig.cell(i, 1), f"(ลงชื่อ) .................................................. ผู้ตรวจรับ\n{name}", align=WD_ALIGN_PARAGRAPH.CENTER)
        if i == 1:
            _set_cell_text(sig.cell(i, 0), "เรียน ผู้ว่าราชการจังหวัดสิงห์บุรี\n- เพื่อทราบผลการตรวจรับ\n\n(นางพิณนภา ศริพันธุ์)\nเจ้าหน้าที่", align=WD_ALIGN_PARAGRAPH.CENTER)
        elif i == 2:
            _set_cell_text(sig.cell(i, 0), "(นายชัชวาล บุญญฤทธิ์)\nหัวหน้าเจ้าหน้าที่", align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            _set_cell_text(sig.cell(i, 0), "", align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, f"หมายเหตุ : เลขที่โครงการ {purchase.project_number or '........................'}", space_after=0)
    _paragraph(doc, f"            เลขคุมสัญญา {purchase.contract_control_number or '........................'}", space_after=0)


def _add_integrity_form(doc, purchase):
    _paragraph(doc, "แบบแสดงความบริสุทธิ์ใจในการจัดซื้อจัดจ้างทุกวิธีของหน่วยงาน", WD_ALIGN_PARAGRAPH.CENTER, True, 18)
    _paragraph(doc, "ในการเปิดเผยข้อมูลความขัดแย้งทางผลประโยชน์", WD_ALIGN_PARAGRAPH.CENTER, True, 18)
    _paragraph(doc, "ของหัวหน้าพัสดุ เจ้าหน้าที่พัสดุ และคณะกรรมการตรวจรับพัสดุ", WD_ALIGN_PARAGRAPH.CENTER, True, 18)
    _paragraph(doc, "-----------------------------------", WD_ALIGN_PARAGRAPH.CENTER)
    persons = [
        ("นายชัชวาลย์ บุญญฤทธิ์", "หัวหน้าเจ้าหน้าที่"),
        ("นางพิณนภา ศริพันธุ์", "เจ้าหน้าที่พัสดุ"),
        ("นางสาวกัญญพัชร ธรกิจการค้า", "คณะกรรมการตรวจรับพัสดุ"),
        ("นางสาวชุลีพร สุขมี", "คณะกรรมการตรวจรับพัสดุ"),
        ("นางสาวกัญญาพัชร เลิศอนันตกูล", "คณะกรรมการตรวจรับพัสดุ"),
    ]
    for name, role in persons:
        _paragraph(doc, f"ข้าพเจ้า {name} ({role})")
    _paragraph(
        doc,
        "ขอให้คำรับรองว่าไม่มีความเกี่ยวข้องหรือมีส่วนได้ส่วนเสียไม่ว่าโดยตรงหรือโดยอ้อม "
        "หรือผลประโยชน์ใด ๆ ที่ก่อให้เกิดความขัดแย้งทางผลประโยชน์กับผู้ขาย ผู้รับจ้าง "
        "ผู้เสนองาน หรือผู้ชนะประมูล หรือผู้มีส่วนเกี่ยวข้องที่เข้ามามีนิติสัมพันธ์ "
        "และวางตัวเป็นกลางในการดำเนินการเกี่ยวกับการพัสดุ ปฏิบัติหน้าที่ด้วยจิตสำนึก "
        "ด้วยความโปร่งใส สามารถให้ผู้เกี่ยวข้องตรวจสอบได้ทุกเวลา และมุ่งประโยชน์ส่วนรวมเป็นสำคัญ",
        first_line=True,
    )
    _paragraph(
        doc,
        "หากปรากฏว่าเกิดความขัดแย้งทางผลประโยชน์ระหว่างข้าพเจ้ากับผู้ขาย ผู้รับจ้าง "
        "ผู้เสนองาน หรือผู้ชนะประมูล หรือผู้มีส่วนเกี่ยวข้องที่เข้ามามีนิติสัมพันธ์ "
        "ข้าพเจ้าจะรายงานให้ทราบโดยทันที",
        first_line=True,
    )
    _paragraph(doc, "")
    for name, role in persons:
        _paragraph(doc, f"ลงนาม .................................................................\n({name})\n({role})", WD_ALIGN_PARAGRAPH.CENTER)
        _paragraph(doc, "")
    _paragraph(doc, f"หมายเหตุ : สำหรับใบสั่งซื้อเลขที่ {purchase.po_number} ลงวันที่ {format_thai_date_full(purchase.document_date)}")


def _add_procurement_pack(doc, purchase):
    """Create the government procurement document pack from the supplied examples."""
    # 1) Purchase request memorandum
    _add_garuda(doc)
    _paragraph(doc, "บันทึกข้อความ", WD_ALIGN_PARAGRAPH.CENTER, True, 20)
    _paragraph(doc, "ส่วนราชการ โรงพยาบาลสิงห์บุรี กลุ่มงานเภสัชกรรม โทร. ๐๓๖๕๒ ๒๕๐๘ ต่อ ๑๑๒๙")
    _paragraph(doc, f"ที่ สห ๐๐๓๓.๒๐๕.๑๒/........................ วันที่ {format_thai_date_full(purchase.document_date)}")
    subject = purchase.lines[0].description if len(purchase.lines) == 1 else f"เวชภัณฑ์มิใช่ยา จำนวน {len(purchase.lines)} รายการ"
    _paragraph(doc, f"เรื่อง รายงานขอซื้อ {subject}")
    _paragraph(doc, "เรียน ผู้ว่าราชการจังหวัดสิงห์บุรี")
    _paragraph(doc, f"ด้วยโรงพยาบาลสิงห์บุรีมีความประสงค์จะซื้อ {subject} โดยวิธีเฉพาะเจาะจง ซึ่งมีรายละเอียดดังต่อไปนี้", first_line=True)
    _paragraph(doc, "๑. เหตุผลความจำเป็นที่ต้องซื้อ\n    ใช้ในการรักษาผู้ป่วย")
    _paragraph(doc, f"๒. รายละเอียดของพัสดุ\n    {subject}")
    _paragraph(doc, f"๓. ราคากลางของพัสดุ เป็นเงิน {purchase.total_amount:,.2f} บาท")
    _paragraph(doc, f"๔. วงเงินที่จะซื้อ {purchase.budget_source or 'เงินบำรุงโรงพยาบาลสิงห์บุรี'} จำนวน {purchase.total_amount:,.2f} บาท ({baht_text(purchase.total_amount)})")
    _paragraph(doc, f"๕. กำหนดเวลาส่งมอบภายใน {purchase.delivery_days} วัน นับถัดจากวันลงนามในสัญญา")
    _paragraph(doc, "๖. ดำเนินการโดยวิธีเฉพาะเจาะจง")
    _paragraph(doc, "๗. หลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอโดยใช้เกณฑ์ราคา")
    _paragraph(doc, "๘. ผู้ตรวจรับพัสดุ\n    ๑. นางสาวกัญญพัชร ธนกิจการค้า ประธานกรรมการ\n    ๒. นางสาวชุลีพร สุขมี กรรมการ\n    ๓. นางสาวกัญญาพัชร เลิศอนันตกูล กรรมการ")
    _paragraph(doc, "จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ", first_line=True)
    _paragraph(doc, "(นางพิณนภา ศริพันธุ์)\nเจ้าหน้าที่", WD_ALIGN_PARAGRAPH.CENTER)

    # 2) Consideration memorandum
    doc.add_page_break()
    _add_garuda(doc)
    _paragraph(doc, "บันทึกข้อความ", WD_ALIGN_PARAGRAPH.CENTER, True, 20)
    _paragraph(doc, f"เรื่อง รายงานผลการพิจารณาและขออนุมัติสั่งซื้อ {subject}")
    _paragraph(doc, "เรียน ผู้ว่าราชการจังหวัดสิงห์บุรี")
    _paragraph(doc, f"ขอรายงานผลการพิจารณาซื้อ {subject} โดยวิธีเฉพาะเจาะจง ดังนี้", first_line=True)
    table = doc.add_table(rows=2, cols=4)
    _apply_table_grid(table)
    headers = ["รายการพิจารณา", "รายชื่อผู้ยื่นข้อเสนอ", "ราคาที่เสนอ*", "ราคาที่ตกลงซื้อหรือจ้าง*"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, True, WD_ALIGN_PARAGRAPH.CENTER)
    vals = [subject, purchase.company.name, f"{purchase.total_amount:,.2f}", f"{purchase.total_amount:,.2f}"]
    for i, v in enumerate(vals):
        _set_cell_text(table.rows[1].cells[i], v, align=WD_ALIGN_PARAGRAPH.CENTER if i != 0 else WD_ALIGN_PARAGRAPH.LEFT)
    _paragraph(doc, "*ราคาที่เสนอและราคาที่ตกลงซื้อหรือจ้างเป็นราคารวมภาษีมูลค่าเพิ่ม ภาษีอื่น ค่าขนส่ง ค่าจดทะเบียน และค่าใช้จ่ายอื่นทั้งปวง", size=14)
    _paragraph(doc, "โรงพยาบาลสิงห์บุรีพิจารณาแล้ว เห็นสมควรจัดซื้อจากผู้เสนอราคาดังกล่าว จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ", first_line=True)
    _paragraph(doc, "(นางพิณนภา ศริพันธุ์)\nเจ้าหน้าที่", WD_ALIGN_PARAGRAPH.CENTER)

    # 3) Winner announcement
    doc.add_page_break()
    _add_garuda(doc)
    _paragraph(doc, "ประกาศจังหวัดสิงห์บุรี", WD_ALIGN_PARAGRAPH.CENTER, True, 20)
    _paragraph(doc, f"เรื่อง ประกาศผู้ชนะการเสนอราคา ซื้อ {subject}", WD_ALIGN_PARAGRAPH.CENTER, True, 18)
    _paragraph(doc, "โดยวิธีเฉพาะเจาะจง", WD_ALIGN_PARAGRAPH.CENTER, True, 18)
    _paragraph(doc, "-----------------------------------", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, f"ตามที่โรงพยาบาลสิงห์บุรีได้มีโครงการซื้อ {subject} โดยวิธีเฉพาะเจาะจงนั้น ผู้ได้รับการคัดเลือก ได้แก่ {purchase.company.name} โดยเสนอราคาเป็นเงินทั้งสิ้น {purchase.total_amount:,.2f} บาท ({baht_text(purchase.total_amount)}) รวมภาษีมูลค่าเพิ่มและค่าใช้จ่ายอื่นทั้งปวง", first_line=True)
    _paragraph(doc, f"ประกาศ ณ วันที่ {format_thai_date_full(purchase.document_date)}", WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "(นายพิรุณ ปิตะหงษ์นันท์)\nผู้อำนวยการโรงพยาบาลสิงห์บุรี ปฏิบัติราชการแทน\nผู้ว่าราชการจังหวัดสิงห์บุรี", WD_ALIGN_PARAGRAPH.CENTER)

def _replace_text_once_in_paragraph(paragraph, old, new):
    """Replace text across runs while preserving surrounding formatting."""
    old = str(old or "")
    new = str(new or "")
    if not old or not paragraph.runs:
        return False

    full_text = "".join(run.text for run in paragraph.runs)
    start_index = full_text.find(old)
    if start_index < 0:
        return False
    end_index = start_index + len(old)

    positions = []
    cursor = 0
    for index, run in enumerate(paragraph.runs):
        positions.append((index, cursor, cursor + len(run.text)))
        cursor += len(run.text)

    first_index = next((i for i, s, e in positions if e > start_index), None)
    last_index = None
    for i, s, e in positions:
        if s < end_index:
            last_index = i

    if first_index is None or last_index is None:
        return False

    first_run = paragraph.runs[first_index]
    last_run = paragraph.runs[last_index]
    first_start = positions[first_index][1]
    last_start = positions[last_index][1]

    prefix = first_run.text[:max(0, start_index - first_start)]
    suffix = last_run.text[max(0, end_index - last_start):]
    first_run.text = prefix + new + suffix

    # ต้อง blank run ที่เหลือ "ก่อน" เรียก _set_run_font เสมอ เพราะถ้า first_run มี
    # เลขปนตัวอักษร _set_run_font จะแตก first_run ออกเป็นหลาย run (เพื่อแยกฟอนต์
    # เลข/ข้อความ) ทำให้ index ของ run ถัดไปใน paragraph.runs เลื่อนหนีจาก
    # first_index/last_index ที่คำนวณไว้ก่อนหน้า ถ้าสลับลำดับจะไป blank run ผิดตัว
    for index in range(first_index + 1, last_index + 1):
        paragraph.runs[index].text = ""

    _set_run_font(first_run, 16)
    return True


def _replace_paragraph_text(paragraph, replacements):
    # Replace each source phrase at most once per paragraph.
    # Using a while-loop can repeat forever when the replacement text
    # still contains the source phrase.
    for old, new in replacements:
        _replace_text_once_in_paragraph(paragraph, old, new)


def _replace_in_document(doc, replacements):
    for paragraph in doc.paragraphs:
        _replace_paragraph_text(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_paragraph_text(paragraph, replacements)


def _thai_document_date(value):
    months = [
        "", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม",
        "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม",
        "พฤศจิกายน", "ธันวาคม",
    ]
    return f"{value.day} {months[value.month]} {value.year + 543}"


def _set_table_value(table, row_index, column_index, value, align=WD_ALIGN_PARAGRAPH.CENTER):
    if row_index < len(table.rows) and column_index < len(table.columns):
        _set_cell_text(table.cell(row_index, column_index), value, align=align, size=16)


def _find_table(doc, required_text):
    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if str(required_text) in table_text:
            return table
    return None


def _replace_exact_xml_texts(doc, source_text, values):
    """Replace exact text nodes, including text inside Word text boxes."""
    pending = [str(value) for value in values]
    if not pending:
        return 0

    replaced = 0
    for text_node in doc.element.xpath(".//w:t"):
        if text_node.text == source_text and pending:
            new_val = pending.pop(0)
            run_element = text_node.getparent()
            if run_element is not None and run_element.tag.endswith("}r"):
                for child in list(run_element):
                    if child.tag in _RUN_CONTENT_TAGS:
                        run_element.remove(child)
                _write_split_text(run_element, new_val)
            else:
                text_node.text = new_val
            replaced += 1
    return replaced


def _fill_calculated_template_values(
    doc,
    subtotal,
    vat,
    current_amount,
    remaining_amount,
):
    # ช่อง "คำนวนมา" สองช่อง:
    # ช่องแรก = ยอดก่อน VAT, ช่องที่สอง = VAT
    _replace_exact_xml_texts(
        doc,
        "คำนวนมา",
        [f"{subtotal:,.2f}", f"{vat:,.2f}"],
    )
    _replace_exact_xml_texts(
        doc,
        "คำนวณมา",
        [f"{subtotal:,.2f}", f"{vat:,.2f}"],
    )

    # ช่อง "คำนวนอัตโนมัติ" สองช่อง:
    # ช่องแรก = ยอดที่จัดหาครั้งนี้
    # ช่องที่สอง = ยอดคงเหลือ
    budget_values = [
        f"{current_amount:,.2f}",
        f"{remaining_amount:,.2f}",
    ]
    replaced = _replace_exact_xml_texts(
        doc,
        "คำนวนอัตโนมัติ",
        budget_values,
    )
    if replaced == 0:
        _replace_exact_xml_texts(
            doc,
            "คำนวณอัตโนมัติ",
            budget_values,
        )



def _arabic_digits(value):
    """Convert Thai numerals to Arabic numerals for seller/contact fields."""
    table = str.maketrans("๐๑๒๓๔๕๖๗๘๙", "0123456789")
    return str(value or "").translate(table)


def _all_standard_paragraphs(doc):
    """Yield normal paragraphs and paragraphs inside normal Word tables."""
    seen = set()

    for paragraph in doc.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    key = id(paragraph._p)
                    if key not in seen:
                        seen.add(key)
                        yield paragraph


def _set_paragraph_compact(paragraph, line_spacing=1.0, before=0, after=0):
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def _set_checkbox_size(paragraph, size=11):
    """Make checkbox glyphs smaller without reducing the surrounding text."""
    for run in list(paragraph.runs):
        if "☐" not in run.text and "□" not in run.text:
            continue

        original = run.text
        # When the run contains only a checkbox, resize it directly.
        if original.strip() in {"☐", "□"}:
            run.font.size = Pt(size)
            continue

        # Split mixed text so only the checkbox is smaller.
        parts = []
        current = ""
        for char in original:
            if char in {"☐", "□"}:
                if current:
                    parts.append(("text", current))
                    current = ""
                parts.append(("box", char))
            else:
                current += char
        if current:
            parts.append(("text", current))

        run.text = ""
        parent = run._r.getparent()
        insert_at = parent.index(run._r) + 1

        for kind, value in parts:
            new_run = paragraph.add_run(value)
            _set_run_font(new_run, size if kind == "box" else 16)
            parent.remove(new_run._r)
            parent.insert(insert_at, new_run._r)
            insert_at += 1


def _replace_first_paragraph_containing(doc, needle, new_text, align=None):
    """Replace a whole standard paragraph containing needle."""
    for paragraph in _all_standard_paragraphs(doc):
        if needle in paragraph.text:
            paragraph.text = ""
            run = paragraph.add_run(new_text)
            _set_run_font(run, 16)
            if align is not None:
                paragraph.alignment = align
            _set_paragraph_compact(paragraph)
            return True
    return False


def _apply_review_layout(doc, purchase, profile, company, thai_date, subtotal, vat, total):
    """Apply the layout corrections marked in the review document."""

    # 1) Headings such as "บันทึกข้อความ" and "ใบสั่งซื้อ" must be centered.
    for paragraph in _all_standard_paragraphs(doc):
        clean = " ".join(paragraph.text.split())

        if clean in {"บันทึกข้อความ", "ใบสั่งซื้อ", "ประกาศจังหวัดสิงห์บุรี"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_compact(paragraph, after=0)

        # Keep ordinary document text tightly spaced like the reviewed copy.
        if clean.startswith(("ส่วนราชการ", "ที่ ", "เรื่อง ", "เรียน ")):
            _set_paragraph_compact(paragraph, after=0)

        # Purchase-order conditions: keep 1-6 compact.
        if clean.startswith(("๑.", "๒.", "๓.", "๔.", "๕.", "๖.")):
            _set_paragraph_compact(paragraph, after=0)

        # Reviewer asked for condition 7 to begin on the following page.
        if clean.startswith("๗.") and "การประเมินผลการปฏิบัติงานของผู้ประกอบการ" in clean:
            paragraph.paragraph_format.page_break_before = True
            _set_paragraph_compact(paragraph, after=0)

        if "หมายเหตุ" in clean:
            _set_paragraph_compact(paragraph, after=0)

        # Acceptance receipt spacing.
        if clean.startswith(("1. ผลการตรวจรับ", "2. ค่าปรับ", "3. การเบิกจ่ายเงิน")):
            _set_paragraph_compact(paragraph, after=0)

        if "☐" in paragraph.text or "□" in paragraph.text:
            _set_paragraph_compact(paragraph, after=0)
            _set_checkbox_size(paragraph, size=11)

    # 2) Announcement date must not remain blank.
    _replace_first_paragraph_containing(
        doc,
        "ประกาศ ณ วันที่",
        f"ประกาศ ณ วันที่ {thai_date}",
        WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 3) Acceptance-receipt date uses Thai long date.
    for paragraph in _all_standard_paragraphs(doc):
        clean = " ".join(paragraph.text.split())
        if clean.startswith("วันที่ ") and "/" in clean:
            paragraph.text = ""
            run = paragraph.add_run(f"วันที่ {thai_date}")
            _set_run_font(run, 16)
            _set_paragraph_compact(paragraph)

        if "ตามใบสั่งซื้อ" in clean and "ลงวันที่" in clean:
            old_short = purchase.document_date.strftime("%d/%m/%Y")
            if old_short in paragraph.text:
                paragraph.text = paragraph.text.replace(old_short, thai_date)
                for run in paragraph.runs:
                    _set_run_font(run, 16)
                _set_paragraph_compact(paragraph)

    # 4) Keep table rows aligned and compact.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    _set_paragraph_compact(paragraph, after=0)

    # 5) Purchase-order numeric rows.
    # These are also handled by XML replacement, but this guarantees normal
    # table cells receive the same calculated values.
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            if "รวมเป็นเงิน" in row_text and "รวมเป็นเงินทั้งสิ้น" not in row_text:
                if len(row.cells) >= 2:
                    _set_cell_text(
                        row.cells[-1],
                        f"{subtotal:,.2f}",
                        align=WD_ALIGN_PARAGRAPH.RIGHT,
                        size=16,
                    )
            elif "ภาษีมูลค่าเพิ่ม" in row_text:
                _set_cell_text(
                    row.cells[-1],
                    f"{vat:,.2f}",
                    align=WD_ALIGN_PARAGRAPH.RIGHT,
                    size=16,
                )
            elif "รวมเป็นเงินทั้งสิ้น" in row_text:
                _set_cell_text(
                    row.cells[-1],
                    f"{total:,.2f}",
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.RIGHT,
                    size=16,
                )



def _xml_paragraphs(doc):
    # A drawing/text-box can contain its own w:p elements inside an outer w:p.
    # Work only on leaf paragraphs so we never collapse several visible
    # text-box lines into one paragraph.
    return doc.element.xpath(".//w:p[not(.//w:p)]")


def _xml_paragraph_text(paragraph_element):
    """อ่านข้อความของ paragraph โดยแปลง <w:br/> เป็น "\\n" และ <w:tab/> เป็น "\\t"
    (เหมือน python-docx Run.text) ไม่ใช่แค่ join <w:t> เฉย ๆ เพราะถ้าไม่แปลง จะทำให้
    ตอนเขียนกลับผ่าน _write_split_text บรรทัดเดิมที่มี <w:br/> ขึ้นบรรทัดใหม่/เว้น
    บรรทัดอยู่แล้วหายไปหมด กลายเป็นข้อความยาวบรรทัดเดียว."""
    parts = []
    for node in paragraph_element.iter():
        tag = node.tag
        if tag == qn("w:t"):
            parts.append(node.text or "")
        elif tag == qn("w:br") or tag == qn("w:cr"):
            parts.append("\n")
        elif tag == qn("w:tab"):
            parts.append("\t")
    return "".join(parts)


# ฟอนต์สัญลักษณ์ที่ไม่มีตัวอักษรไทย/เลขไทย ต้องข้ามไว้ ไม่งั้น glyph หาย
_SYMBOL_FONT_NAMES = {"Segoe UI Symbol", "Wingdings", "Wingdings 2", "Wingdings 3", "Webdings", "Symbol"}


def _normalize_document_fonts(doc):
    """บังคับให้ทุก run ที่เหลือใน template (label ที่ไม่ได้ผ่าน _set_run_font
    หรือ _set_xml_paragraph_text) ใช้ชุดฟอนต์ตัวเลข/ข้อความเดียวกันทั้งเอกสาร
    รวมถึงแยกเลข 0-9 ออกจากตัวอักษรที่ปนอยู่ใน run เดียวกันด้วย.

    ใช้ python-docx Run.text (ไม่ใช่ join w:t เอง) เพื่ออ่าน <w:br/>/<w:tab/> ที่แทรก
    อยู่ใน run เป็น \\n/\\t ให้ถูกต้อง แล้ว _write_split_text จะคืนกลับเป็น
    <w:br/>/<w:tab/> ตอนเขียน ไม่งั้นบรรทัดจะพังหายเวลาเปิดใน Word."""
    for paragraph_element in _xml_paragraphs(doc):
        for run in paragraph_element.findall(qn("w:r")):
            if not any(run.findall(tag) for tag in _RUN_CONTENT_TAGS):
                continue
            rpr = run.find(qn("w:rPr"))
            rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
            current_ascii = rfonts.get(qn("w:ascii")) if rfonts is not None else None
            if current_ascii in _SYMBOL_FONT_NAMES:
                continue
            text = DocxRun(run, None).text
            _write_split_text(run, text)


def _set_xml_paragraph_text(paragraph_element, value, size=16):
    text_nodes = paragraph_element.xpath(".//w:t")
    if text_nodes:
        run = text_nodes[0].getparent()

        # _write_split_text เขียน "value" (ซึ่งรวมเนื้อหา+ตัวขึ้นบรรทัดใหม่ทั้งพารากราฟ
        # แล้ว เพราะ value มาจาก _xml_paragraph_text ที่อ่านทั้งพารากราฟ) ลง run เดียว
        # แล้วแทรก run ใหม่ต่อจากตำแหน่งนี้ ดังนั้น run/w:t อื่นที่เหลือในพารากราฟ
        # ซ้ำซ้อนกับเนื้อหาที่เขียนใหม่แล้วทั้งหมด ต้องลบทิ้ง ไม่ใช่แค่ blank ข้อความ
        # ไม่งั้น <w:br/> เดิมที่ยังค้างอยู่จะไปโผล่เป็นบรรทัดว่างซ้ำท้ายพารากราฟ
        run_parent = run.getparent()
        seen_run = False
        for child in list(run_parent):
            if child is run:
                seen_run = True
                continue
            if not seen_run:
                continue
            if child.tag == qn("w:r") or child.tag == qn("w:hyperlink"):
                run_parent.remove(child)

        _write_split_text(run, str(value))
        return

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rpr.append(sz)
    run.append(rpr)

    _write_split_text(run, str(value))
    paragraph_element.append(run)


def _replace_all_xml_paragraphs(doc, replacements):
    for paragraph_element in _xml_paragraphs(doc):
        original = _xml_paragraph_text(paragraph_element)
        if not original:
            continue
        updated = original
        for old, new in replacements:
            old = str(old or "")
            if old and old in updated:
                updated = updated.replace(old, str(new or ""))
        if updated != original:
            _set_xml_paragraph_text(paragraph_element, updated, 16)


def _xml_set_alignment(paragraph_element, value):
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph_element.insert(0, ppr)

    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(qn("w:val"), value)


def _xml_set_spacing(paragraph_element, before="0", after="0", line="240"):
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph_element.insert(0, ppr)

    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)

    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")


def _xml_disable_snap_to_grid(paragraph_element):
    """Turn off "Snap to grid" (w:snapToGrid) for one paragraph.

    The document sections use a document grid (``w:docGrid w:linePitch``)
    for consistent printed line spacing. When that is on, Word still
    vertically snaps every line -- including a line holding only a picture
    -- to the next multiple of the grid's line pitch, no matter what
    explicit ``w:spacing`` the paragraph carries. A picture whose height
    isn't an exact multiple of that pitch then renders with a visible dead
    gap underneath it in Word even though the paragraph spacing is already
    zero and other renderers (which don't emulate this quirk) show no gap
    at all. Only the paragraphs the master template's own author already
    fixed this way skip the gap; every other Garuda/heading paragraph must
    get the same explicit override.
    """
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph_element.insert(0, ppr)

    snap = ppr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        ppr.insert(0, snap)
    snap.set(qn("w:val"), "0")


def _xml_page_break_before(paragraph_element):
    ppr = paragraph_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph_element.insert(0, ppr)

    node = ppr.find(qn("w:pageBreakBefore"))
    if node is None:
        node = OxmlElement("w:pageBreakBefore")
        ppr.append(node)
    node.set(qn("w:val"), "1")


def _fill_budget_xml_block(doc, allocated, previously_used, current_amount, remaining):
    paragraphs = _xml_paragraphs(doc)
    texts = [_xml_paragraph_text(p) for p in paragraphs]

    for index, value in enumerate(texts):
        if "ยอดที่ได้รับจัดสรร" not in value:
            continue

        value_start = index + 4
        values = [
            f"{allocated:,.2f}",
            f"{previously_used:,.2f}",
            f"{current_amount:,.2f}",
            f"{remaining:,.2f}",
        ]

        for offset, output_value in enumerate(values):
            target = value_start + offset
            if target < len(paragraphs):
                _set_xml_paragraph_text(paragraphs[target], output_value, 16)
        break


def _fill_purchase_order_calculation_xml(doc, subtotal, vat, total):
    paragraphs = _xml_paragraphs(doc)
    texts = [_xml_paragraph_text(p) for p in paragraphs]

    for index, value in enumerate(texts):
        if value.strip() != "รวมเป็นเงิน":
            continue

        if index + 1 < len(paragraphs):
            _set_xml_paragraph_text(paragraphs[index + 1], f"{subtotal:,.2f}", 16)

        for vat_index in range(index + 1, min(index + 8, len(paragraphs))):
            if texts[vat_index].strip() != "ภาษีมูลค่าเพิ่ม":
                continue

            if vat_index + 1 < len(paragraphs):
                _set_xml_paragraph_text(paragraphs[vat_index + 1], f"{vat:,.2f}", 16)

            for total_index in range(vat_index + 1, min(vat_index + 8, len(paragraphs))):
                if texts[total_index].strip() == "รวมเป็นเงินทั้งสิ้น":
                    if total_index - 1 >= 0:
                        _set_xml_paragraph_text(
                            paragraphs[total_index - 1],
                            f"({baht_text(total)})",
                            16,
                        )
                    if total_index + 1 < len(paragraphs):
                        _set_xml_paragraph_text(
                            paragraphs[total_index + 1],
                            f"{total:,.2f}",
                            16,
                        )
                    return



def _fix_spec_date_xml(doc, purchase, thai_date):
    """Fix duplicated/old date text on the specification memorandum."""
    day = str(purchase.document_date.day)

    for paragraph_element in _xml_paragraphs(doc):
        value = _xml_paragraph_text(paragraph_element)
        clean = " ".join(value.split())

        # Handles examples such as:
        # "วันที่ 31 31 กรกฎาคม 2569"
        # "วันที่ 25 25 กรกฎาคม 2569"
        duplicated = f"{day} {thai_date}"
        if "วันที่" in clean and duplicated in value:
            _set_xml_paragraph_text(
                paragraph_element,
                value.replace(duplicated, thai_date),
                16,
            )

        # Exact date-only field beside the government letter number.
        clean = " ".join(_xml_paragraph_text(paragraph_element).split())
        if clean.startswith("วันที่") and (
            "/" in clean
            or clean.endswith("กรกฎาคม 2569")
            or clean.endswith("กรกฎาคม ๒๕๖๙")
        ):
            # Do not touch signature blank date lines.
            if "........" not in clean:
                _set_xml_paragraph_text(
                    paragraph_element,
                    f"วันที่      {thai_date}",
                    16,
                )


def _fix_integrity_note_xml(doc, purchase, thai_date):
    """Preserve the full integrity-form note instead of replacing the whole line."""
    expected = (
        f"หมายเหตุ : สำหรับใบสั่งซื้อเลขที่ {purchase.po_number} "
        f"ลงวันที่ {thai_date}"
    )

    for paragraph_element in _xml_paragraphs(doc):
        value = _xml_paragraph_text(paragraph_element)
        clean = " ".join(value.split())

        if (
            "หมายเหตุ" in clean
            and "สำหรับใบสั่งซื้อเลขที่" in clean
        ):
            _set_xml_paragraph_text(paragraph_element, expected, 16)
            return


def _fix_original_po_product_row_xml(doc, purchase):
    """Update the original purchase-order product row stored in text boxes."""
    if not purchase.lines:
        return

    line = purchase.lines[0]
    paragraphs = _xml_paragraphs(doc)
    texts = [_xml_paragraph_text(p).strip() for p in paragraphs]

    # Use the first matching item in the original PO region only.
    # A nearby sequence is: item description, quantity, unit, unit price, amount.
    for index, value in enumerate(texts):
        if value != line.description:
            continue

        # Avoid the later generated/spec tables by looking for the original
        # Thai numeral "๑" shortly before this row.
        before = " ".join(texts[max(0, index - 8):index])
        if "๑" not in before and "ลำดับ" not in before:
            continue

        nearby = list(range(index + 1, min(index + 12, len(paragraphs))))
        assigned_quantity = False
        assigned_unit = False
        assigned_price = False
        assigned_amount = False

        for j in nearby:
            current = texts[j]
            normalized = current.replace(",", "").strip()

            if not current:
                continue

            if not assigned_quantity:
                # Template quantity is a simple number such as 10.
                try:
                    Decimal(normalized)
                    _set_xml_paragraph_text(
                        paragraphs[j],
                        f"{line.quantity:g}",
                        16,
                    )
                    assigned_quantity = True
                    continue
                except Exception:
                    pass

            if assigned_quantity and not assigned_unit:
                if current.upper() == line.unit.name.upper() or current.isalpha():
                    _set_xml_paragraph_text(
                        paragraphs[j],
                        line.unit.name,
                        16,
                    )
                    assigned_unit = True
                    continue

            if assigned_quantity and assigned_unit and not assigned_price:
                try:
                    Decimal(normalized)
                    _set_xml_paragraph_text(
                        paragraphs[j],
                        f"{line.unit_price:,.2f}",
                        16,
                    )
                    assigned_price = True
                    continue
                except Exception:
                    pass

            if assigned_price and not assigned_amount:
                try:
                    Decimal(normalized)
                    _set_xml_paragraph_text(
                        paragraphs[j],
                        f"{line.amount:,.2f}",
                        16,
                    )
                    assigned_amount = True
                    break
                except Exception:
                    pass

        if assigned_quantity:
            return


def _fix_original_po_xml(
    doc,
    purchase,
    company,
    thai_date,
    phone_arabic,
    tax_arabic,
    account_arabic,
):
    paragraphs = _xml_paragraphs(doc)

    for i, paragraph_element in enumerate(paragraphs):
        value = _xml_paragraph_text(paragraph_element)
        clean = " ".join(value.split())

        if "โทรศัพท์" in clean and "แปลงเป็นเลขอาราบิก" in clean:
            _set_xml_paragraph_text(paragraph_element, f"โทรศัพท์   {phone_arabic}", 16)

        if "เลขประจำตัวผู้เสียภาษี" in clean:
            _set_xml_paragraph_text(
                paragraph_element,
                f"เลขประจำตัวผู้เสียภาษี   {tax_arabic}",
                16,
            )

        if "เลขที่บัญชีเงินฝากธนาคาร" in clean:
            _set_xml_paragraph_text(
                paragraph_element,
                f"เลขที่บัญชีเงินฝากธนาคาร   {account_arabic}",
                16,
            )

        if (
            clean.startswith("ใบสั่งซื้อเลขที่")
            and "หมายเหตุ" not in clean
            and "สำหรับใบสั่งซื้อ" not in clean
        ):
            _set_xml_paragraph_text(
                paragraph_element,
                f"ใบสั่งซื้อเลขที่  {purchase.po_number}",
                16,
            )

            for j in range(i + 1, min(i + 5, len(paragraphs))):
                next_text = " ".join(_xml_paragraph_text(paragraphs[j]).split())
                if next_text == "วันที่":
                    _set_xml_paragraph_text(paragraphs[j], f"วันที่  {thai_date}", 16)
                    break

        if "ได้เสนอราคา" in clean and "โรงพยาบาลสิงห์บุรี" in clean:
            updated = value
            for old in (
                "บริษัท แปซิฟิค เฮลธ์แคร์ (ไทยแลนด์) จำกัด",
                "บริษัท แปซิฟิค เฮลธ์แคร์(ไทยแลนด์) จำกัด",
            ):
                updated = updated.replace(old, company.name)
            if updated != value:
                _set_xml_paragraph_text(paragraph_element, updated, 16)



def _fix_request_budget_amount_xml(doc, purchase):
    total = to_decimal(purchase.total_amount)
    total_text = f"{total:,.2f}"

    for paragraph_element in _xml_paragraphs(doc):
        value = _xml_paragraph_text(paragraph_element)
        clean = " ".join(value.split())

        if (
            "วงเงินที่จะซื้อ" not in clean
            and "เงินนอกงบประมาณจาก" in clean
            and "จำนวน" in clean
            and "บาท" in clean
        ):
            # Replace any money token immediately after "จำนวน".
            import re
            updated = re.sub(
                r"(จำนวน\\s+)[0-9,]+(?:\\.\\d{2})?(\\s*บาท)",
                rf"\\g<1>{total_text}\\2",
                value,
                count=1,
            )
            if updated != value:
                _set_xml_paragraph_text(paragraph_element, updated, 16)


def _fix_all_word_xml_layout(doc, purchase, thai_date):
    for paragraph_element in _xml_paragraphs(doc):
        value = _xml_paragraph_text(paragraph_element)
        clean = " ".join(value.split())
        if not clean:
            continue

        if clean in {"บันทึกข้อความ", "ใบสั่งซื้อ", "ประกาศจังหวัดสิงห์บุรี"}:
            _xml_set_alignment(paragraph_element, "center")

        if clean.startswith((
            "ส่วนราชการ",
            "ที่ ",
            "เรื่อง ",
            "เรียน ",
            "๑.",
            "๒.",
            "๓.",
            "๔.",
            "๕.",
            "๖.",
        )):
            _xml_set_spacing(paragraph_element, before="0", after="0", line="240")

        if clean.startswith("๗.") and "การประเมินผลการปฏิบัติงาน" in clean:
            _xml_page_break_before(paragraph_element)
            _xml_set_spacing(paragraph_element, before="0", after="0", line="240")

        if "หมายเหตุ" in clean:
            _xml_set_spacing(paragraph_element, before="0", after="0", line="240")

        duplicated = f"{purchase.document_date.day} {thai_date}"
        if duplicated in value:
            _set_xml_paragraph_text(
                paragraph_element,
                value.replace(duplicated, thai_date),
                16,
            )



def _fix_corrupted_sara_am(doc):
    """แก้ "˚า" (U+02DA ring above + U+0E32 สระอา ที่ถูกพิมพ์/แปลงแยกกันผิด) ให้กลับเป็น
    "ำ" (U+0E33 สระอำ ตัวเดียว) ทั่วทั้งเอกสาร ปัญหานี้ฝังอยู่ใน purchase_master.docx
    เอง (131 จุด) มาตั้งแต่ก่อนแก้เรื่องฟอนต์ใด ๆ ไม่ใช่บั๊กจากการแยกฟอนต์เลข/ข้อความ."""
    for text_node in doc.element.body.iter(qn("w:t")):
        if text_node.text and "˚า" in text_node.text:
            text_node.text = text_node.text.replace("˚า", "ำ")


def _center_all_garuda_paragraphs(doc):
    """
    Center embedded Garuda/picture objects.

    Word pictures can be either:
      1) inline pictures -> paragraph alignment controls them
      2) floating wp:anchor pictures -> paragraph alignment does not move them

    Therefore both cases are handled here.
    """
    # Inline pictures and drawing paragraphs.
    for paragraph in doc.paragraphs:
        has_picture = bool(
            paragraph._element.xpath(".//w:drawing")
            or paragraph._element.xpath(".//w:pict")
        )
        if has_picture:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # The document uses a document grid (see _xml_disable_snap_to_grid),
            # which snaps every line -- including a picture-only line -- to
            # the grid's line pitch in Word regardless of the paragraph's own
            # zero spacing. Without this, Word (unlike other renderers) shows
            # a dead gap under the emblem even though the file looks correct
            # everywhere else.
            _xml_disable_snap_to_grid(paragraph._element)

    # Floating pictures: convert to inline so they are centered by paragraph alignment
    for anchor in doc.element.xpath(".//wp:anchor"):
        inline = OxmlElement("wp:inline")
        for k, v in anchor.attrib.items():
            if k not in ['behindDoc', 'locked', 'layoutInCell', 'allowOverlap', 'simplePos', 'relativeHeight']:
                inline.set(k, v)
        
        for child in list(anchor):
            if child.tag not in [qn('wp:simplePos'), qn('wp:positionH'), qn('wp:positionV'), qn('wp:wrapNone'), qn('wp:wrapSquare'), qn('wp:wrapTight'), qn('wp:wrapThrough'), qn('wp:wrapTopAndBottom')]:
                inline.append(child)
                
        parent = anchor.getparent()
        parent.replace(anchor, inline)

    # Inline pictures inside table cells/text boxes may not appear in
    # doc.paragraphs, so center every leaf paragraph containing a drawing.
    #
    # These pictures (the Garuda embedded directly inside the master
    # template's memo table) never got an explicit paragraph spacing
    # override, so they were still falling back to the template's default
    # "Normal" style spacing (10pt after + 1.15 line height). That leftover
    # gap showed up right under the Garuda emblem on almost every page --
    # the picture paragraph must be flush against whatever follows it.
    for p in doc.element.xpath(".//w:p[not(.//w:p)]"):
        if not (p.xpath(".//w:drawing") or p.xpath(".//w:pict")):
            continue

        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)

        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")

        _xml_set_spacing(p, before="0", after="0", line="240")
        _xml_disable_snap_to_grid(p)



def _remove_blank_line_above_memo(doc):
    from docx.shared import Pt
    from docx.text.paragraph import Paragraph
    for p in doc.element.xpath(".//w:p"):
        text = "".join(t.text for t in p.xpath(".//w:t") if t.text)
        if "บันทึก" in text and "ข้อความ" in text:
            parent = p.getparent()
            
            # Remove space before/after for the memo paragraph itself
            para_obj = Paragraph(p, parent)
            para_obj.paragraph_format.space_before = Pt(0)
            para_obj.paragraph_format.space_after = Pt(0)
            
            paragraphs = parent.xpath("./w:p")
            p_idx = paragraphs.index(p)
            if p_idx > 0:
                prev = paragraphs[p_idx-1]
                prev_text = "".join(t.text for t in prev.xpath(".//w:t") if t.text)
                if not prev_text.strip():
                    parent.remove(prev)
            elif parent.tag.endswith("tc"):
                tc = parent
                tr = tc.getparent()
                tbl = tr.getparent()
                tr_idx = tbl.index(tr)
                
                prev_tr = None
                for i in range(tr_idx-1, -1, -1):
                    if tbl[i].tag.endswith("tr"):
                        prev_tr = tbl[i]
                        break
                        
                if prev_tr is not None:
                    tcs = prev_tr.xpath("./w:tc")
                    if len(tcs) > 0:
                        prev_tc = tcs[0]
                        prev_paragraphs = prev_tc.xpath("./w:p")
                        if prev_paragraphs:
                            last_p = prev_paragraphs[-1]
                            last_p_text = "".join(t.text for t in last_p.xpath(".//w:t") if t.text)
                            if not last_p_text.strip():
                                prev_tc.remove(last_p)
                                # After removing, the new last paragraph is the one with the Garuda
                                new_prev_paragraphs = prev_tc.xpath("./w:p")
                                if new_prev_paragraphs:
                                    garuda_p = new_prev_paragraphs[-1]
                                    garuda_para_obj = Paragraph(garuda_p, prev_tc)
                                    garuda_para_obj.paragraph_format.space_before = Pt(0)
                                    garuda_para_obj.paragraph_format.space_after = Pt(0)


def _center_all_garuda_paragraphs(doc):
    """
    Center embedded Garuda/picture objects.

    Word pictures can be either:
      1) inline pictures -> paragraph alignment controls them
      2) floating wp:anchor pictures -> paragraph alignment does not move them

    Therefore both cases are handled here.
    """
    # Inline pictures and drawing paragraphs.
    for paragraph in doc.paragraphs:
        has_picture = bool(
            paragraph._element.xpath(".//w:drawing")
            or paragraph._element.xpath(".//w:pict")
        )
        if has_picture:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # The document uses a document grid (see _xml_disable_snap_to_grid),
            # which snaps every line -- including a picture-only line -- to
            # the grid's line pitch in Word regardless of the paragraph's own
            # zero spacing. Without this, Word (unlike other renderers) shows
            # a dead gap under the emblem even though the file looks correct
            # everywhere else.
            _xml_disable_snap_to_grid(paragraph._element)

    # Floating pictures: convert to inline so they are centered by paragraph alignment
    for anchor in doc.element.xpath(".//wp:anchor"):
        inline = OxmlElement("wp:inline")
        for k, v in anchor.attrib.items():
            if k not in ['behindDoc', 'locked', 'layoutInCell', 'allowOverlap', 'simplePos', 'relativeHeight']:
                inline.set(k, v)
        
        for child in list(anchor):
            if child.tag not in [qn('wp:simplePos'), qn('wp:positionH'), qn('wp:positionV'), qn('wp:wrapNone'), qn('wp:wrapSquare'), qn('wp:wrapTight'), qn('wp:wrapThrough'), qn('wp:wrapTopAndBottom')]:
                inline.append(child)
                
        parent = anchor.getparent()
        parent.replace(anchor, inline)

    # Inline pictures inside table cells/text boxes may not appear in
    # doc.paragraphs, so center every leaf paragraph containing a drawing.
    #
    # These pictures (the Garuda embedded directly inside the master
    # template's memo table) never got an explicit paragraph spacing
    # override, so they were still falling back to the template's default
    # "Normal" style spacing (10pt after + 1.15 line height). That leftover
    # gap showed up right under the Garuda emblem on almost every page --
    # the picture paragraph must be flush against whatever follows it.
    for p in doc.element.xpath(".//w:p[not(.//w:p)]"):
        if not (p.xpath(".//w:drawing") or p.xpath(".//w:pict")):
            continue

        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)

        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")

        _xml_set_spacing(p, before="0", after="0", line="240")
        _xml_disable_snap_to_grid(p)



def _remove_blank_line_above_memo(doc):
    from docx.shared import Pt
    for p in doc.element.xpath(".//w:p"):
        text = "".join(t.text for t in p.xpath(".//w:t") if t.text)
        if "บันทึก" in text and "ข้อความ" in text:
            parent = p.getparent()
            paragraphs = parent.xpath("./w:p")
            p_idx = paragraphs.index(p)
            
            # Remove space before/after for the memo paragraph itself
            pPr = p.find(".//w:pPr")
            if pPr is not None:
                spacing = pPr.find(".//w:spacing")
                if spacing is not None:
                    spacing.set("before", "0")
                    spacing.set("after", "0")
            
            if p_idx > 0:
                prev = paragraphs[p_idx-1]
                prev_text = "".join(t.text for t in prev.xpath(".//w:t") if t.text)
                if not prev_text.strip():
                    parent.remove(prev)
            elif parent.tag.endswith("tc"):
                tc = parent
                tr = tc.getparent()
                tbl = tr.getparent()
                tr_idx = tbl.index(tr)
                
                prev_tr = None
                for i in range(tr_idx-1, -1, -1):
                    if tbl[i].tag.endswith("tr"):
                        prev_tr = tbl[i]
                        break
                        
                if prev_tr is not None:
                    tcs = prev_tr.xpath("./w:tc")
                    if len(tcs) > 0:
                        prev_tc = tcs[0]
                        prev_paragraphs = prev_tc.xpath("./w:p")
                        if prev_paragraphs:
                            last_p = prev_paragraphs[-1]
                            last_p_text = "".join(t.text for t in last_p.xpath(".//w:t") if t.text)
                            if not last_p_text.strip():
                                prev_tc.remove(last_p)
                            # Remove space after for the Garuda paragraph (now the new last_p)
                            new_prev_paragraphs = prev_tc.xpath("./w:p")
                            if new_prev_paragraphs:
                                garuda_p = new_prev_paragraphs[-1]
                                garuda_pPr = garuda_p.find(".//w:pPr")
                                if garuda_pPr is not None:
                                    spacing = garuda_pPr.find(".//w:spacing")
                                    if spacing is not None:
                                        spacing.set("before", "0")
                                        spacing.set("after", "0")

def _center_all_garuda_paragraphs(doc):
    """
    Center embedded Garuda/picture objects.

    Word pictures can be either:
      1) inline pictures -> paragraph alignment controls them
      2) floating wp:anchor pictures -> paragraph alignment does not move them

    Therefore both cases are handled here.
    """
    # Inline pictures and drawing paragraphs.
    for paragraph in doc.paragraphs:
        has_picture = bool(
            paragraph._element.xpath(".//w:drawing")
            or paragraph._element.xpath(".//w:pict")
        )
        if has_picture:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # The document uses a document grid (see _xml_disable_snap_to_grid),
            # which snaps every line -- including a picture-only line -- to
            # the grid's line pitch in Word regardless of the paragraph's own
            # zero spacing. Without this, Word (unlike other renderers) shows
            # a dead gap under the emblem even though the file looks correct
            # everywhere else.
            _xml_disable_snap_to_grid(paragraph._element)

    # Floating pictures: convert to inline so they are centered by paragraph alignment
    for anchor in doc.element.xpath(".//wp:anchor"):
        inline = OxmlElement("wp:inline")
        for k, v in anchor.attrib.items():
            if k not in ['behindDoc', 'locked', 'layoutInCell', 'allowOverlap', 'simplePos', 'relativeHeight']:
                inline.set(k, v)
        
        for child in list(anchor):
            if child.tag not in [qn('wp:simplePos'), qn('wp:positionH'), qn('wp:positionV'), qn('wp:wrapNone'), qn('wp:wrapSquare'), qn('wp:wrapTight'), qn('wp:wrapThrough'), qn('wp:wrapTopAndBottom')]:
                inline.append(child)
                
        parent = anchor.getparent()
        parent.replace(anchor, inline)

    # Inline pictures inside table cells/text boxes may not appear in
    # doc.paragraphs, so center every leaf paragraph containing a drawing.
    #
    # These pictures (the Garuda embedded directly inside the master
    # template's memo table) never got an explicit paragraph spacing
    # override, so they were still falling back to the template's default
    # "Normal" style spacing (10pt after + 1.15 line height). That leftover
    # gap showed up right under the Garuda emblem on almost every page --
    # the picture paragraph must be flush against whatever follows it.
    for p in doc.element.xpath(".//w:p[not(.//w:p)]"):
        if not (p.xpath(".//w:drawing") or p.xpath(".//w:pict")):
            continue

        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)

        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")

        _xml_set_spacing(p, before="0", after="0", line="240")
        _xml_disable_snap_to_grid(p)



def _remove_blank_line_above_memo(doc):
    for p in doc.element.xpath(".//w:p"):
        text = "".join(t.text for t in p.xpath(".//w:t") if t.text)
        if "บันทึก" in text and "ข้อความ" in text:
            parent = p.getparent()
            paragraphs = parent.xpath("./w:p")
            p_idx = paragraphs.index(p)
            # Some memo sections stack more than one leftover blank/
            # whitespace-only paragraph directly above the heading (e.g. a
            # line of nbsp characters *and* a line of plain spaces). Keep
            # removing the immediate previous paragraph as long as it is
            # blank, instead of stopping after the first one, otherwise the
            # Garuda emblem still ends up with a visible gap under it.
            while p_idx > 0:
                prev = paragraphs[p_idx - 1]
                prev_text = "".join(t.text for t in prev.xpath(".//w:t") if t.text)
                if prev_text.strip():
                    break
                parent.remove(prev)
                del paragraphs[p_idx - 1]
                p_idx -= 1

            if p_idx == 0 and parent.tag.endswith("tc"):
                tc = parent
                tr = tc.getparent()
                tbl = tr.getparent()
                tr_idx = tbl.index(tr)

                prev_tr = None
                for i in range(tr_idx-1, -1, -1):
                    if tbl[i].tag.endswith("tr"):
                        prev_tr = tbl[i]
                        break

                if prev_tr is not None:
                    tcs = prev_tr.xpath("./w:tc")
                    if len(tcs) > 0:
                        prev_tc = tcs[0]
                        while True:
                            prev_paragraphs = prev_tc.xpath("./w:p")
                            if not prev_paragraphs:
                                break
                            last_p = prev_paragraphs[-1]
                            last_p_text = "".join(t.text for t in last_p.xpath(".//w:t") if t.text)
                            has_picture = bool(
                                last_p.xpath(".//w:drawing") or last_p.xpath(".//w:pict")
                            )
                            if last_p_text.strip() or has_picture:
                                break
                            prev_tc.remove(last_p)


def _remove_dead_default_spacing_paragraphs(doc):
    """Delete leftover blank paragraphs still using Word's default "Normal"
    style spacing (10pt after + 1.15 line height) instead of the document's
    compact spacing.

    The master template was originally spaced out by pressing Enter several
    times between each "บันทึกข้อความ" section instead of using a real page
    break. Every other paragraph in the template has an explicit compact
    ``w:spacing`` override, but these blank leftovers never got one, so each
    still reserves a full default-styled line. Nine of them sit right above
    the Garuda emblem/heading of the "รายงานผลการพิจารณา" memo, which is the
    dead space the customer saw pushing that memo (and its table) below the
    page and onto a third page. ``_apply_customer_procurement_layout`` below
    already puts an explicit page-break-before on that memo's Garuda
    paragraph, so removing this leftover padding cannot merge it back onto
    the previous page -- it only pulls the memo's own content back up so it
    fits on a single page again.
    """
    for paragraph_element in list(doc.element.body):
        if paragraph_element.tag != qn("w:p"):
            continue
        if _xml_paragraph_text(paragraph_element).strip():
            continue
        if paragraph_element.xpath(".//w:drawing") or paragraph_element.xpath(".//w:pict"):
            continue

        p_pr = paragraph_element.find(qn("w:pPr"))
        spacing = p_pr.find(qn("w:spacing")) if p_pr is not None else None
        if spacing is not None:
            # Already has an explicit (intentionally compact) spacing value.
            continue

        doc.element.body.remove(paragraph_element)


def _apply_customer_procurement_layout(doc):
    """Apply the latest customer-marked layout corrections to the Word master.

    - A later ``บันทึกข้อความ`` starts at the top of a fresh page so the
      Garuda/heading and the table below are not pushed down by signatures from
      the previous page.
    - Remove only blank paragraphs between the Garuda and ``บันทึกข้อความ``.
    - Keep the Garuda and memo heading spacing compact without rebuilding the
      rest of the government form.
    """
    for table_index, table in enumerate(doc.tables):
        rows = list(table.rows)
        for row_index, row in enumerate(rows):
            row_text = " ".join(
                "".join(t.text or "" for t in cell._tc.xpath(".//w:t"))
                for cell in row.cells
            )
            if "บันทึกข้อความ" not in row_text:
                continue

            # Memo heading: no extra paragraph spacing.
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    xml_text = "".join(
                        t.text or "" for t in paragraph._element.xpath(".//w:t")
                    )
                    if "บันทึกข้อความ" in xml_text:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0
                        # Keep this heading's line off the document grid too
                        # (see _xml_disable_snap_to_grid) -- when the memo's
                        # Garuda and heading sit in separate table rows
                        # (unlike the first memo, where both share one row)
                        # a grid-snapped heading line still renders with a
                        # gap above it in Word even with spacing at zero.
                        _xml_disable_snap_to_grid(paragraph._element)

            if row_index == 0:
                continue

            garuda_row = rows[row_index - 1]
            for cell in garuda_row.cells:
                # Remove trailing blank paragraphs after the picture.
                while len(cell.paragraphs) > 1:
                    paragraph = cell.paragraphs[-1]
                    has_picture = bool(
                        paragraph._element.xpath(".//w:drawing")
                        or paragraph._element.xpath(".//w:pict")
                    )
                    if paragraph.text.strip() or has_picture:
                        break
                    paragraph._element.getparent().remove(paragraph._element)

                for paragraph in cell.paragraphs:
                    has_picture = bool(
                        paragraph._element.xpath(".//w:drawing")
                        or paragraph._element.xpath(".//w:pict")
                    )
                    if not has_picture:
                        continue
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0

                    # The first memo is already at the beginning of the report.
                    # Later memo sections must start on a new page so the
                    # Garuda sits near the top and the following table has room.
                    if table_index > 0:
                        paragraph.paragraph_format.page_break_before = True


def _build_procurement_pack(purchase):
    """สร้าง "ชุดรายงานจัดซื้อ" (รายงานขอซื้อ + รายงานผลการพิจารณา + ประกาศผู้ชนะ)
    จาก procurement_pack_master.docx — เวอร์ชันสะอาดที่ไม่มีปัญหา "˚า"/character
    scaling เหมือน purchase_master.docx เดิม."""
    template_path = Path(__file__).resolve().parent / "templates" / "word" / "procurement_pack_master.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"ไม่พบไฟล์ต้นแบบ Word: {template_path}")

    doc = Document(str(template_path))

    profile = (
        purchase.government_profile
        or GovernmentProfile.query.filter_by(active=True).first()
        or GovernmentProfile()
    )

    lines = list(purchase.lines)
    item_count = len(lines)
    total = to_decimal(purchase.total_amount)
    procurement_type = purchase.procurement_type or "เวชภัณฑ์มิใช่ยา"
    procurement_label = f"{procurement_type} จำนวน {item_count} รายการ"
    thai_date = _thai_document_date(purchase.document_date)

    budget_allocated = to_decimal(purchase.budget_allocated)
    budget_used = to_decimal(purchase.budget_previously_used)
    budget_this_time = total
    budget_remaining = (budget_allocated - budget_used - budget_this_time).quantize(Decimal("0.01"))

    company = purchase.company

    # เอกสารแยกบรรทัด "ผู้อำนวยการ...ปฏิบัติราชการแทน" กับ "ผู้ว่าราชการ..." ออกจากกัน
    # (บรรทัดที่ 2 เยื้องซ้าย) แต่ profile.approver_position (ค่า default) เก็บสอง
    # บรรทัดนี้รวมเป็นสตริงเดียว ถ้าเอา approver_position ทั้งก้อนไปแทนที่บรรทัดแรก
    # เฉย ๆ แล้วมี replacement ของ "ผู้ว่าราชการจังหวัดสิงห์บุรี" (บรรทัด 2) แยกอีกจุด
    # จะได้ "ผู้ว่าราชการจังหวัดสิงห์บุรี" โผล่ซ้ำ 2 ครั้งติดกัน (เคยเจอปัญหานี้มาแล้ว)
    # จึงต้องตัดข้อความ recipient ที่ซ้อนท้าย approver_position ออกก่อน
    approver_position_line1 = (profile.approver_position or "").rstrip()
    recipient_suffix = (profile.recipient or "").strip()
    if recipient_suffix and approver_position_line1.endswith(recipient_suffix):
        approver_position_line1 = approver_position_line1[: -len(recipient_suffix)].rstrip()

    replacements = [
        ("โรงพยาบาลสิงห์บุรี กลุ่มงานเภสัชกรรมโทร. ๐ ๓๖๕๒ ๒๕๐๘ ต่อ ๑๑๒๙", profile.department),
        ("โรงพยาบาลสิงห์บุรี กลุ่มงานเภสัชกรรม โทร. ๐ ๓๖๕๒ ๒๕๐๘ ต่อ ๑๑๒๙", profile.department),
        ("สห ๐๐๓๓.๒๐๕.๑๒/", profile.letter_prefix),
        ("กรกฎาคม ๒๕๖๙", thai_date),

        ("รายงานขอซื้อเวชภัณฑ์มิใช่ยา จำนวน 1 รายการ", f"รายงานขอซื้อ{procurement_label}"),
        (
            "รายงานผลการพิจารณาและขออนุมัติสั่งซื้อเวชภัณฑ์มิใช่ยา จำนวน 1 รายการ",
            f"รายงานผลการพิจารณาและขออนุมัติสั่งซื้อ{procurement_label}",
        ),
        (
            "เรื่อง ประกาศผู้ชนะการเสนอราคา เวชภัณฑ์มิใช่ยา จำนวน 1 รายการ",
            f"เรื่อง ประกาศผู้ชนะการเสนอราคา {procurement_label}",
        ),
        ("เวชภัณฑ์มิใช่ยา จำนวน 1 รายการ", procurement_label),

        ("บริษัท พี.เอ็น.โปรดักส์ นครสวรรค์ จำกัด", company.name),
        ("ขายส่ง,ขายปลีก,ให้บริการ", company.business_type or "-"),

        # เรียงจากยาว/เจาะจงที่สุดไปสั้นที่สุดเสมอ เพราะ "-29,010.00" มี "10.00" เป็น
        # substring อยู่ข้างใน ถ้าแทน "10.00" ก่อนจะไปกิน "-29,010.00" ให้เพี้ยน
        ("-29,010.00", f"{budget_remaining:,.2f}"),
        ("30,000.00", f"{total:,.2f}"),
        ("สามหมื่นบาทถ้วน", baht_text(total)),
        ("1,000.00", f"{budget_allocated:,.2f}"),
        ("10.00", f"{budget_used:,.2f}"),

        ("นางสาวกัญญพัชร ธนกิจการค้า", profile.inspector1_name),
        ("นางสาวชุลีพร สุขมี", profile.inspector2_name),
        ("นางสาวกัญญาพัชร เลิศอนันตกูล", profile.inspector3_name),
        ("เจ้าพนักงานเภสัชกรรมชำนาญงาน", profile.inspector2_position),
        ("เจ้าพนักงานเภสัชกรรม", profile.inspector3_position),
        ("เภสัชกร", profile.inspector1_position),

        ("นางพิณนภา ศริพันธุ์", profile.officer_name),
        ("นายชัชวาลย์ บุญญฤทธิ์", profile.chief_name),
        ("เภสัชกรชำนาญการพิเศษ", profile.chief_position),
        ("นายพิรุณ\xa0 ปิตะหงษ์นันท์", profile.approver_name),
        ("นายพิรุณ ปิตะหงษ์นันท์", profile.approver_name),
        ("ผู้อำนวยการโรงพยาบาลสิงห์บุรี ปฏิบัติราชการแทน", approver_position_line1),
        ("ผู้ว่ารายการจังหวัดสิงห์บุรี", profile.recipient),
        ("ผู้ว่าราชการจังหวัดสิงห์บุรี", profile.recipient),
    ]
    _replace_in_document(doc, replacements)
    _replace_all_xml_paragraphs(doc, replacements)

    _replace_first_paragraph_containing(
        doc,
        "ประกาศ ณ วันที่",
        f"ประกาศ ณ วันที่ {thai_date}",
        WD_ALIGN_PARAGRAPH.CENTER,
    )

    _remove_blank_line_above_memo(doc)
    _remove_dead_default_spacing_paragraphs(doc)
    _center_all_garuda_paragraphs(doc)
    _apply_customer_procurement_layout(doc)
    _normalize_document_fonts(doc)

    return doc


def _build_exact_procurement_template(purchase):
    """Fill the manually corrected Word master without adding duplicate pages."""
    template_path = Path(__file__).resolve().parent / "templates" / "word" / "purchase_master.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"ไม่พบไฟล์ต้นแบบ Word: {template_path}")

    doc = Document(str(template_path))

    # ไฟล์ต้นแบบมีข้อความ "ำ" (U+0E33 สระอำ) เพี้ยนเป็น "˚า" (U+02DA + U+0E32,
    # ring above + สระอา แยกกัน) อยู่ 131 จุดทั่วเอกสาร มาตั้งแต่ต้น (ไม่เกี่ยวกับ
    # ฟอนต์ใด ๆ) ทำให้คำอย่าง "จำนวน" ขึ้นเป็น "จ˚านวน" — แก้ตรงนี้ก่อนแทนที่ค่าอื่น ๆ
    _fix_corrupted_sara_am(doc)

    # Center Garuda images already embedded in the Word master.
    _center_all_garuda_paragraphs(doc)
    profile = (
        purchase.government_profile
        or GovernmentProfile.query.filter_by(active=True).first()
        or GovernmentProfile()
    )

    lines = list(purchase.lines)
    item_count = len(lines)
    total = to_decimal(purchase.total_amount)
    subtotal, vat = _vat_values(purchase)
    procurement_type = purchase.procurement_type or "เวชภัณฑ์มิใช่ยา"
    procurement_label = f"{procurement_type} จำนวน {item_count} รายการ"
    thai_date = _thai_document_date(purchase.document_date)
    short_date = purchase.document_date.strftime("%d/%m/%Y")

    budget_allocated = to_decimal(purchase.budget_allocated)
    budget_used = to_decimal(purchase.budget_previously_used)
    budget_this_time = total
    budget_remaining = (budget_allocated - budget_used - budget_this_time).quantize(Decimal("0.01"))

    company = purchase.company

    # The form stores the whole budget-source phrase.  The original template
    # already contains the prefix in some places, so normalize it once to
    # prevent "เงินนอกงบประมาณจาก" from appearing twice.
    budget_source_full = (
        purchase.budget_source
        or "เงินนอกงบประมาณจาก เงินบำรุงโรงพยาบาลสิงห์บุรี ปี ๒๕๖๙"
    ).strip()
    budget_source_without_prefix = budget_source_full
    prefix = "เงินนอกงบประมาณจาก "
    if budget_source_without_prefix.startswith(prefix):
        budget_source_without_prefix = budget_source_without_prefix[len(prefix):].strip()

    company_phone_arabic = _arabic_digits(company.phone or "-")
    company_tax_arabic = _arabic_digits(company.tax_id or "-")
    company_account_arabic = _arabic_digits(company.account_no or "-")

    replacements = [
        ("โรงพยาบาลสิงห์บุรี กลุ่มงานเภสัชกรรมโทร. ๐ ๓๖๕๒ ๒๕๐๘ ต่อ ๑๑๒๙", profile.department),
        ("โรงพยาบาลสิงห์บุรี กลุ่มงานเภสัชกรรม โทร. ๐ ๓๖๕๒ ๒๕๐๘ ต่อ ๑๑๒๙", profile.department),
        ("โรงพยาบาลสิงห์บุรี   กลุ่มงานเภสัชกรรม โทร. 03652 2508 ต่อ 1129", profile.department),
        ("สห ๐๐๓๓.๒๐๕.๑๒/", profile.letter_prefix),
        ("กรกฎาคม   2569", thai_date),
        ("กรกฎาคม ๒๕๖๙", thai_date),
        ("กรกฎาคม 2569", thai_date),

        ("เวชภัณฑ์มิใช่ยา จำนวน 10 รายการ", procurement_label),
        ("เวชภัณฑ์มิใช่ยา จำวน 10 รายการ", procurement_label),
        ("เวชภัณฑ์มิใช่ยา จำนวน 2 รายการ", procurement_label),
        ("เวชภัณฑ์มิใช่ยา จำนวน2 รายการ", procurement_label),
        ("เวชภัณฑ์มิใช่ยา จำนวน 1 รายการ", procurement_label),
        ("เวชภัณฑ์มิใช่ยา จำวน 1 รายการ", procurement_label),

        ("บริษัท พี.เอ็น.โปรดักส์ นครสวรรค์ จำกัด", company.name),
        ("บริษัท พี.เอ็น.โปรดักส์ นครสวรรค์จำกัด", company.name),
        ("บริษัท แปซิฟิค เฮลธ์แคร์ (ไทยแลนด์) จำกัด", company.name),
        ("บริษัท แปซิฟิค เฮลธ์แคร์(ไทยแลนด์) จำกัด", company.name),
        ("๐๕๖๒๒๒๓๑๒", company_phone_arabic),
        ("๐๖๐๕๕๒๒๐๐๐๗๙๗", company_tax_arabic),
        ("๖๒๘๑๒๘๐๙๑๑", company_account_arabic),
        ("๐๕๖๒๒๒๑๑๒", company_phone_arabic),
        ("๐๖๐๕๕๒๒๐๐๐๗๙๗", company_tax_arabic),
        ("ธนาคารกรุงไทยจำกัด (มหาชน)", company.bank_name or "-"),
        ("ปากน้ำโพ", company.bank_branch or "-"),
        ("๖๒๘๑๒๘๐๙๘๑", company_account_arabic),
        ("พี.เอ็น.โปรดักส์ นครสวรรค์ หจก.", company.account_name or "-"),

        ("PO-๖๙-0๒00๗๘", purchase.po_number),
        ("PO-๖๙-๐๒00821", purchase.po_number),
        ("PO-๖๙-๐๒๐๐802", purchase.po_number),
        ("25/07/2026", short_date),
        ("69079275357", purchase.project_number or "........................"),
        ("690714258286", purchase.contract_control_number or "........................"),

        ("30,000.00", f"{total:,.2f}"),
        ("หนึ่งหมื่นสี่พันบาทถ้วน", baht_text(total)),
        ("14,000.00", f"{total:,.2f}"),
        ("30000", f"{total:,.2f}"),
        ("สามหมื่นบาทถ้วน", baht_text(total)),
        ("28,037.38", f"{subtotal:,.2f}"),
        ("1,962.62", f"{vat:,.2f}"),

        ("1,000.00", f"{budget_allocated:,.2f}"),
        ("1000.00", f"{budget_allocated:,.2f}"),

        ("ใช้ในการรักษาผู้ป่วย", purchase.necessity_reason or "ใช้ในการรักษาผู้ป่วย"),
        (
            "เงินนอกงบประมาณจาก เงินนอกงบประมาณจาก เงินบำรุงโรงพยาบาลสิงห์บุรี ปี ๒๕๖๙",
            budget_source_full,
        ),
        (
            "เงินนอกงบประมาณจาก เงินบำรุงโรงพยาบาลสิงห์บุรี ปี ๒๕๖๙",
            budget_source_full,
        ),
        (
            "เงินบำรุงโรงพยาบาลสิงห์บุรี ปี ๒๕๖๙",
            budget_source_without_prefix,
        ),
        ("โรงพยาบาลสิงห์บุรี ๙๑๗/๓", purchase.delivery_place or "โรงพยาบาลสิงห์บุรี ๙๑๗/๓"),

        ("นางพิณนภา ศริพันธุ์", profile.officer_name),
        ("นายชัชวาลย์ บุญญฤทธิ์", profile.chief_name),
        ("นายชัชวาล บุญญฤทธิ์", profile.chief_name),
        ("นายพิรุณ ปิตะหงษ์นันท์", profile.approver_name),
        ("นางสาวกัญญพัชร ธนกิจการค้า", profile.inspector1_name),
        ("นางสาวชุลีพร สุขมี", profile.inspector2_name),
        ("นางสาวกัญญาพัชร เลิศอนันตกูล", profile.inspector3_name),
        ("นางสาวนลินี เครือทิวา", profile.specifier_name),
        ("ผู้ว่ารายการจังหวัดสิงห์บุรี", "ผู้ว่าราชการจังหวัดสิงห์บุรี"),
        ("ธนาคา รกรุงไทยจำกัด", "ธนาคารกรุงไทยจำกัด"),
        ("(แปลงเป็นเลขอาราบิก)", ""),
        ("25 25 กรกฎาคม 2569", thai_date),
    ]
    _replace_in_document(doc, replacements)
    _replace_all_xml_paragraphs(doc, replacements)

    # เติมค่าที่คำนวณแล้วลงใน Text Box ของไฟล์ Word
    _fill_calculated_template_values(
        doc,
        subtotal=subtotal,
        vat=vat,
        current_amount=budget_this_time,
        remaining_amount=budget_remaining,
    )

    # Update all product tables while retaining the template's layout.
    for table in doc.tables:
        header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        if not ("รายการ" in header and "จำนวน" in header and ("หน่วยละ" in header or "ราคาต่อหน่วย" in header)):
            continue

        product_row_count = min(6, max(0, len(table.rows) - 1))
        for index in range(product_row_count):
            line = lines[index] if index < item_count else None
            row_index = index + 1

            if len(table.columns) == 5:
                values = [
                    str(index + 1),
                    line.description if line else "",
                    f"{line.quantity:g} {line.unit.name}" if line else "",
                    f"{line.unit_price:,.2f}" if line else "",
                    f"{line.amount:,.2f}" if line else "",
                ]
            else:
                values = [
                    str(index + 1),
                    line.description if line else "",
                    f"{line.quantity:g}" if line else "",
                    line.unit.name if line else "",
                    f"{line.unit_price:,.2f}" if line else "",
                    f"{line.amount:,.2f}" if line else "",
                ]

            for column_index, value in enumerate(values):
                if column_index < len(table.columns):
                    align = WD_ALIGN_PARAGRAPH.LEFT if column_index == 1 else WD_ALIGN_PARAGRAPH.CENTER
                    _set_table_value(table, row_index, column_index, value, align)

        for row_index, row in enumerate(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)
            if "รวมเป็นเงินทั้งสิ้น" in row_text:
                _set_table_value(table, row_index, len(table.columns) - 1, f"{total:,.2f}", WD_ALIGN_PARAGRAPH.RIGHT)
                if len(table.columns) >= 5:
                    _set_table_value(table, row_index, 1, f"({baht_text(total)})", WD_ALIGN_PARAGRAPH.CENTER)

    # Replace seller information in the purchase-order box.
    seller_table = _find_table(doc, "ผู้ขาย")
    if seller_table is not None and seller_table.rows:
        left_text = (
            f"ผู้ขาย {company.name}\n"
            f"ที่อยู่ {company.address or '-'}\n"
            f"โทรศัพท์ {company_phone_arabic}\n"
            f"เลขประจำตัวผู้เสียภาษี {company_tax_arabic}\n"
            f"เลขที่บัญชีเงินฝากธนาคาร {company_account_arabic}\n"
            f"ชื่อบัญชี {company.account_name or '-'}\n"
            f"ธนาคาร {company.bank_name or '-'}"
            + (f" สาขา {company.bank_branch}" if company.bank_branch else "")
        )
        right_text = (
            f"เลขที่ {purchase.po_number}\n"
            f"วันที่ {thai_date}\n\n"
            "ส่วนราชการ โรงพยาบาลสิงห์บุรี\n"
            "ที่อยู่ ๙๑๗/๓ ตำบลบางพุทรา อำเภอเมืองสิงห์บุรี จังหวัดสิงห์บุรี ๑๖๐๐๐\n"
            "โทรศัพท์ ๐๓๖-๕๒๒๕๐๗"
        )
        _set_cell_text(seller_table.cell(0, 0), left_text, size=16)
        _format_po_meta_cell(seller_table.cell(0, 0))
        if len(seller_table.columns) > 1:
            _set_cell_text(seller_table.cell(0, 1), right_text, size=16)
            _format_po_meta_cell(seller_table.cell(0, 1))

    # Fill the budget table with calculated values.
    budget_table = _find_table(doc, "ยอดที่ได้รับ")
    if budget_table is not None and len(budget_table.rows) >= 2:
        values = [budget_allocated, budget_used, budget_this_time, budget_remaining]
        for column_index, value in enumerate(values):
            if column_index < len(budget_table.columns):
                _set_table_value(
                    budget_table,
                    1,
                    column_index,
                    f"{value:,.2f}",
                    WD_ALIGN_PARAGRAPH.CENTER,
                )

    _fill_budget_xml_block(
        doc,
        allocated=budget_allocated,
        previously_used=budget_used,
        current_amount=budget_this_time,
        remaining=budget_remaining,
    )
    _fill_purchase_order_calculation_xml(
        doc,
        subtotal=subtotal,
        vat=vat,
        total=total,
    )
    _fix_original_po_xml(
        doc,
        purchase=purchase,
        company=company,
        thai_date=thai_date,
        phone_arabic=company_phone_arabic,
        tax_arabic=company_tax_arabic,
        account_arabic=company_account_arabic,
    )
    _fix_original_po_product_row_xml(
        doc,
        purchase=purchase,
    )
    _fix_request_budget_amount_xml(
        doc,
        purchase=purchase,
    )
    _fix_spec_date_xml(
        doc,
        purchase=purchase,
        thai_date=thai_date,
    )
    _fix_integrity_note_xml(
        doc,
        purchase=purchase,
        thai_date=thai_date,
    )
    _fix_all_word_xml_layout(
        doc,
        purchase=purchase,
        thai_date=thai_date,
    )

    _apply_review_layout(
        doc,
        purchase=purchase,
        profile=profile,
        company=company,
        thai_date=thai_date,
        subtotal=subtotal,
        vat=vat,
        total=total,
    )

    _normalize_document_fonts(doc)

    return doc



def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set Word table cell borders with OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge_name, edge_value in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        if edge_value is None:
            continue

        tag = f"w:{edge_name}"
        edge = tcBorders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tcBorders.append(edge)

        if edge_value is False:
            edge.set(qn("w:val"), "nil")
        else:
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), str(edge_value))
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "000000")


def _set_table_row_height(row, cm_value):
    trPr = row._tr.get_or_add_trPr()
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    trHeight.set(qn("w:val"), str(int(Cm(cm_value))))
    trHeight.set(qn("w:hRule"), "atLeast")


def _set_cell_width(cell, cm_value):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(Cm(cm_value))))
    tcW.set(qn("w:type"), "dxa")


def _write_spec_table_cell(
    cell,
    text="",
    bold=False,
    size=16,
    align=WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(str(text or ""))
    _set_run_font(run, size=size, bold=bold)
    return paragraph


def _set_table_grid_widths(table, widths_cm):
    """Set w:tblGrid explicitly so Word cannot expand the table columns."""
    tbl_grid = table._tbl.tblGrid

    for child in list(tbl_grid):
        tbl_grid.remove(child)

    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(width_cm))))
        tbl_grid.append(grid_col)


def _set_compact_table_cell_margins(table, margin_cm=0.04):
    """Reduce default Word cell padding, which otherwise makes the table wider."""
    tblPr = table._tbl.tblPr
    cell_mar = tblPr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tblPr.append(cell_mar)

    twips = str(int(Cm(margin_cm)))
    for side in ("top", "left", "bottom", "right"):
        node = cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_mar.append(node)
        node.set(qn("w:w"), twips)
        node.set(qn("w:type"), "dxa")







def _add_final_spec_excel_page(doc, purchase):
    """
    หน้าสุดท้าย: สร้างด้วย Word table จริง
    อ้างอิงรูปแบบกรอบตารางจากไฟล์ตัวอย่างที่แก้กรอบแล้ว

    - ตราครุฑอยู่กึ่งกลางหน้า
    - ข้อความนอกตารางทุกบรรทัดอยู่กึ่งกลาง
    - ตารางหลักมีกรอบครบทุกช่อง
    - ตารางงบประมาณมีกรอบครบทุกช่อง
    - ตารางทั้งสองอยู่กึ่งกลางหน้า
    - ไม่ใช้ Excel / รูปภาพแทนตาราง
    """

    doc.add_page_break()

    lines = list(purchase.lines)
    total = to_decimal(purchase.total_amount)

    budget_allocated = to_decimal(purchase.budget_allocated)
    budget_used = to_decimal(purchase.budget_previously_used)
    budget_this_time = total
    budget_remaining = (
        budget_allocated - budget_used - budget_this_time
    ).quantize(Decimal("0.01"))

    # ==========================================================
    # Helpers
    # ==========================================================
    def _force_centered_table_widths(table, widths_cm):
        """
        บังคับความกว้างจริง + จัดกึ่งกลาง
        Word Online จะไม่ดันตารางออกซ้าย/ขวาเอง
        """
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        tbl_pr = table._tbl.tblPr

        # fixed layout
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

        # center table
        tbl_jc = tbl_pr.first_child_found_in("w:jc")
        if tbl_jc is None:
            tbl_jc = OxmlElement("w:jc")
            tbl_pr.append(tbl_jc)
        tbl_jc.set(qn("w:val"), "center")

        # remove any old indent that could push table away from center
        old_ind = tbl_pr.first_child_found_in("w:tblInd")
        if old_ind is not None:
            tbl_pr.remove(old_ind)

        # exact table width
        total_width_cm = sum(widths_cm)
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(int(Cm(total_width_cm))))

        # exact grid widths
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)

        for width_cm in widths_cm:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(int(Cm(width_cm))))
            grid.append(grid_col)

        # exact widths on every cell
        for row in table.rows:
            for col_index, width_cm in enumerate(widths_cm):
                if col_index >= len(row.cells):
                    continue

                cell = row.cells[col_index]
                tc_pr = cell._tc.get_or_add_tcPr()

                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)

                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(int(Cm(width_cm))))

    def _force_table_borders(table, size=8):
        """
        กรอบดำแบบตัวอย่าง:
        - รอบนอก
        - เส้นคั่นทุกแถว
        - เส้นคั่นทุกคอลัมน์
        """
        tbl_pr = table._tbl.tblPr
        tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")

        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)

        for edge in (
            "top", "left", "bottom", "right",
            "insideH", "insideV",
        ):
            node = tbl_borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                tbl_borders.append(node)

            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), str(size))
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), "000000")

        # บังคับระดับ cell ด้วย เพื่อให้ Word Online แสดงเส้นแน่นอน
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                borders = tc_pr.first_child_found_in("w:tcBorders")

                if borders is None:
                    borders = OxmlElement("w:tcBorders")
                    tc_pr.append(borders)

                for edge in ("top", "left", "bottom", "right"):
                    node = borders.find(qn(f"w:{edge}"))
                    if node is None:
                        node = OxmlElement(f"w:{edge}")
                        borders.append(node)

                    node.set(qn("w:val"), "single")
                    node.set(qn("w:sz"), str(size))
                    node.set(qn("w:space"), "0")
                    node.set(qn("w:color"), "000000")

    def _center_cell_text(cell, value="", bold=False, size=13):
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        run = p.add_run(str(value or ""))
        _set_run_font(run, size=size, bold=bold)

    def _center_paragraph(text="", size=16, bold=False, after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.0

        run = p.add_run(str(text or ""))
        _set_run_font(run, size=size, bold=bold)
        return p

    # ==========================================================
    # ตราครุฑ — กลางหน้า
    # ==========================================================
    garuda_path = (
        Path(__file__).resolve().parent
        / "static"
        / "img"
        / "garuda.png"
    )

    if garuda_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

        p.add_run().add_picture(
            str(garuda_path),
            width=Cm(1.65),
        )

    # ==========================================================
    # ข้อความนอกตาราง — กลางทั้งหมด
    # ==========================================================
    _center_paragraph(
        "รายละเอียดคุณลักษณะเฉพาะ",
        size=18,
        bold=True,
        after=0,
    )

    _center_paragraph(
        f"{purchase.procurement_type or 'เวชภัณฑ์มิใช่ยา'} จำนวน {len(lines)} รายการ",
        size=16,
        bold=True,
        after=4,
    )

    # ==========================================================
    # ตารางหลัก — แบบกรอบตัวอย่าง
    # ความกว้างรวม 17.2 cm เพื่อให้อยู่ใน A4 แน่นอน
    # ==========================================================
    widths = [
        1.10,  # ลำดับที่
        3.60,  # รายการ
        5.60,  # คุณลักษณะเฉพาะ
        2.50,  # จำนวนที่ต้องการ
        2.00,  # หน่วยนับ
        2.40,  # หมายเหตุ
    ]

    table = doc.add_table(rows=1, cols=6)
    _force_centered_table_widths(table, widths)

    headers = [
        "ลำดับที่",
        "รายการ",
        "คุณลักษณะเฉพาะ",
        "จำนวนที่ต้องการ",
        "หน่วยนับ",
        "หมายเหตุ",
    ]

    for i, header in enumerate(headers):
        _center_cell_text(
            table.rows[0].cells[i],
            header,
            bold=True,
            size=13,
        )

    # รายการ
    for index, line in enumerate(lines, start=1):
        cells = table.add_row().cells

        specification = (
            line.specification
            or getattr(line.item, "specification", "")
            or ""
        )

        values = [
            str(index),
            line.description,
            specification,
            f"{line.quantity:g}",
            line.unit.name,
            "",
        ]

        for i, value in enumerate(values):
            _center_cell_text(
                cells[i],
                value,
                bold=False,
                size=12.5,
            )

    # แถวรวม
    total_cells = table.add_row().cells

    total_label = total_cells[0].merge(total_cells[2])
    _center_cell_text(
        total_label,
        "รวม",
        bold=True,
        size=13,
    )

    total_qty = sum(
        (to_decimal(line.quantity) for line in lines),
        Decimal("0.00"),
    )

    _center_cell_text(
        total_cells[3],
        f"{total_qty:g}",
        bold=True,
        size=13,
    )

    total_unit = (
        lines[0].unit.name
        if len(lines) == 1
        else ""
    )

    _center_cell_text(
        total_cells[4],
        total_unit,
        bold=True,
        size=13,
    )

    _center_cell_text(
        total_cells[5],
        "",
        bold=False,
        size=13,
    )

    # สำคัญ: หลัง add_row + merge เสร็จ ต้องบังคับ width/grid อีกครั้ง
    _force_centered_table_widths(table, widths)
    _force_table_borders(table, size=8)

    # ช่องว่างกึ่งกลาง
    _center_paragraph("", size=8, bold=False, after=2)

    # ==========================================================
    # ตารางงบประมาณ — แบบกรอบตัวอย่าง
    # กว้างเท่าตารางหลัก = 17.2 cm
    # ==========================================================
    budget_table = doc.add_table(rows=2, cols=4)

    budget_widths = [
        4.30,
        4.30,
        4.30,
        4.30,
    ]

    _force_centered_table_widths(
        budget_table,
        budget_widths,
    )

    budget_headers = [
        "ยอดที่ได้รับจัดสรร",
        "ยอดที่จัดหาแล้ว",
        "ยอดที่จัดหาครั้งนี้",
        "ยอดคงเหลือ",
    ]

    budget_values = [
        budget_allocated,
        budget_used,
        budget_this_time,
        budget_remaining,
    ]

    for i in range(4):
        _center_cell_text(
            budget_table.rows[0].cells[i],
            budget_headers[i],
            bold=True,
            size=13,
        )

        _center_cell_text(
            budget_table.rows[1].cells[i],
            f"{budget_values[i]:,.2f}",
            bold=False,
            size=13,
        )

    _force_centered_table_widths(
        budget_table,
        budget_widths,
    )
    _force_table_borders(
        budget_table,
        size=8,
    )

    # ==========================================================
    # หมายเหตุ — อยู่นอกตารางและกลางทั้งหมด
    # ==========================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(2.2)
    p.paragraph_format.right_indent = Cm(2.2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    note_text = (
        "*ราคาที่เสนอ และราคาที่ตกลงซื้อหรือจ้าง "
        "เป็นราคารวมภาษีมูลค่าเพิ่มและภาษีอื่น "
        "ค่าขนส่ง ค่าจดทะเบียน และค่าใช้จ่ายอื่นๆ ทั้งปวง"
    )

    note_run = p.add_run(note_text)
    _set_run_font(note_run, size=12.5, bold=False)

    return table


def _remove_trailing_empty_paragraphs(doc):
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        # Make sure we don't delete a page break
        has_break = False
        for run in doc.paragraphs[-1].runs:
            if run._element.xpath(".//w:br[@w:type='page']"):
                has_break = True
                break
        if has_break:
            break
        
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)


def _add_spec_details(doc, purchase):
    # No Garuda emblem on this page: it is a plain attachment/schedule
    # (รายละเอียดขอบเขตและคุณลักษณะ) rather than a formal บันทึกข้อความ or
    # ประกาศ, and the customer confirmed the letterhead only belongs on
    # the pages they marked, not this one.

    # Add header text
    _paragraph(doc, "รายละเอียดขอบเขตและคุณลักษณะเวชภัณฑ์มิใช่ยาที่จะซื้อหรือจ้าง", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    _paragraph(doc, "เวชภัณฑ์มิใช่ยา", WD_ALIGN_PARAGRAPH.CENTER, True, 16)
    _paragraph(doc, "โรงพยาบาลสิงห์บุรี", WD_ALIGN_PARAGRAPH.CENTER, True, 16)

    # Add table
    table = doc.add_table(rows=1, cols=6)
    _apply_table_grid(table)
    table.autofit = False
    
    # Define widths (approximate from excel)
    widths = [Cm(1.2), Cm(10.0), Cm(2.0), Cm(1.5), Cm(2.0), Cm(2.3)]
    for i in range(6):
        table.columns[i].width = widths[i]
        table.rows[0].cells[i].width = widths[i]

    # Header Row
    headers = ["ลำดับ", "รายการ", "จำนวน", "", "หน่วยละ", "เป็นเงิน"]
    for i, value in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], value, True, WD_ALIGN_PARAGRAPH.CENTER, size=16)

    # Merge cell 2 and 3 for 'จำนวน' header
    table.rows[0].cells[2].merge(table.rows[0].cells[3])

    # Items
    for line in purchase.lines:
        cells = table.add_row().cells
        values = [
            line.line_no, 
            line.description, 
            f"{line.quantity:g}", 
            line.unit.name, 
            f"{line.unit_price:,.2f}", 
            f"{line.amount:,.2f}"
        ]
        alignments = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER
        ]
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            _set_cell_text(cells[i], value, False, alignments[i], size=16)

    # Empty row for spacing
    empty_cells = table.add_row().cells
    for i in range(6):
        empty_cells[i].width = widths[i]
        _set_cell_text(empty_cells[i], "", False, WD_ALIGN_PARAGRAPH.LEFT, size=16)

    # Bullet points
    bullets = [
        " -ใช้เพื่อประกอบการรักษาพยาบาลผู้ป่วย",
        " -เวชภัณฑ์มิใช่ยาที่ใช้ในการรักษาพยาบาลผู้ป่วยที่ขึ้นทะเบียนกับ อย.ประเทศไทย",
        " -ผลิตภัณฑ์ต้องเป็นของใหม่และไม่เคยใช้งานมาก่อน",
        " -ผลิตภัณฑ์ต้องมีอายุการใช้งานอย่างน้อย 1 ปี นับตั้งแต่วันส่งมอบ"
    ]
    for bullet in bullets:
        row = table.add_row()
        for i in range(6):
            row.cells[i].width = widths[i]
        # Merge columns 1 to 5
        row.cells[1].merge(row.cells[5])
        _set_cell_text(row.cells[1], bullet, False, WD_ALIGN_PARAGRAPH.LEFT, size=16)

    # Empty rows for padding (4 rows)
    for _ in range(4):
        empty_cells = table.add_row().cells
        for i in range(6):
            empty_cells[i].width = widths[i]
            _set_cell_text(empty_cells[i], "", False, WD_ALIGN_PARAGRAPH.LEFT, size=16)

    # Total Row
    total_cells = table.add_row().cells
    for i in range(6):
        total_cells[i].width = widths[i]
    total_cells[1].merge(total_cells[2])
    _set_cell_text(total_cells[1], f"({baht_text(purchase.total_amount)})", True, WD_ALIGN_PARAGRAPH.CENTER, size=16)
    
    total_cells[3].merge(total_cells[4])
    _set_cell_text(total_cells[3], "รวมเป็นเงินทั้งสิ้น", True, WD_ALIGN_PARAGRAPH.CENTER, size=16)
    
    _set_cell_text(total_cells[5], f"{purchase.total_amount:,.2f}", True, WD_ALIGN_PARAGRAPH.CENTER, size=16)

    # Footer text
    _paragraph(doc, "")
    _paragraph(doc, "โดยวิธีซื้อครั้งหลังสุดภายในระยะเวลา ๒ ปีงบประมาณ พิจารณาคัดเลือกข้อเสนอโดยใช้หลักเกณฑ์ราคา", WD_ALIGN_PARAGRAPH.LEFT, False, 16)
    _paragraph(doc, "")
    _paragraph(doc, "")
    
    # Signature
    _paragraph(doc, "                         (ลงชื่อ)...........................................................ผู้กำหนดรายละเอียดคุณลักษณะเฉพาะ", WD_ALIGN_PARAGRAPH.LEFT, False, 16)
    
    spec_name = purchase.government_profile.specifier_name if purchase.government_profile else "............................................"
    spec_pos = purchase.government_profile.specifier_position if purchase.government_profile else "............................................"
    
    _paragraph(doc, f"                                      ({spec_name})", WD_ALIGN_PARAGRAPH.LEFT, False, 16)
    _paragraph(doc, f"                                 {spec_pos}", WD_ALIGN_PARAGRAPH.LEFT, False, 16)



def _make_word_fully_editable(doc):
    """Remove Word editing restrictions and object locks before export.

    This is intentionally applied to every generated .docx so users can edit
    normal text, table cells and text inside drawing/text-box objects in Word.
    It also cleans up restrictions that may be introduced by a future master
    template (document protection, permission ranges or content-control locks).
    """
    root = doc.element

    # 1) Document-level protection / write protection.
    settings = doc.settings._element
    for tag in ("w:documentProtection", "w:writeProtection"):
        for node in list(settings.findall(qn(tag))):
            settings.remove(node)

    # 2) Permission-range restrictions inside the document body.
    for xpath in (".//w:permStart", ".//w:permEnd"):
        for node in list(root.xpath(xpath)):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)

    # 3) Content controls: remove any lock flags.  Keep the controls themselves
    # because unwrapping them can alter carefully tuned government-form layout.
    for sdt_pr in root.xpath(".//w:sdtPr"):
        for lock in list(sdt_pr.findall(qn("w:lock"))):
            sdt_pr.remove(lock)

    # 4) Drawing / text-box locks.  Word may treat anchored shapes and grouped
    # shapes as locked even when the text around them remains editable.
    for node in root.xpath(".//wp:anchor | .//wp:inline"):
        node.set("locked", "0")

    # DrawingML locking elements can contain flags such as noSelect/noTextEdit.
    # Remove them completely so text boxes and grouped form labels are editable.
    # Remove lock elements by local tag name instead of namespace-specific
    # XPath, because older VML templates can use namespace prefixes that are
    # not registered in python-docx's XPath helper.
    lock_element_names = {"graphicFrameLocks", "spLocks", "grpSpLocks", "picLocks", "lock"}
    for node in list(root.iter()):
        local_name = node.tag.rsplit("}", 1)[-1] if isinstance(node.tag, str) else ""
        if local_name in lock_element_names:
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)

    # 5) Some templates carry lock attributes directly on drawing nodes.
    # Explicitly clear the common forms without changing positions/sizing.
    for node in root.iter():
        for attr in list(node.attrib):
            local = attr.rsplit("}", 1)[-1]
            if local in {"locked", "noSelect", "noTextEdit", "noMove", "noResize"}:
                node.attrib.pop(attr, None)


def _build_word(purchase, form_type):
    doc = Document()
    _configure_document(doc)

    if form_type in {"po_single", "po_multiple", "po"}:
        _add_purchase_order(doc, purchase)
    elif form_type == "spec":
        _add_spec(doc, purchase)
    elif form_type in {"receipt_single", "receipt_two", "receipt_multiple", "receipt"}:
        _add_acceptance_receipt(doc, purchase)
    elif form_type == "integrity":
        _add_integrity_form(doc, purchase)
    elif form_type == "procurement_pack":
        doc = _build_procurement_pack(purchase)
    elif form_type == "all":
        # เรียง 8, 2, 3, 7, 6 ตามที่ผู้ใช้ต้องการ
        # 8. ชุดรายงานจัดซื้อ (procurement_pack)
        doc = _build_procurement_pack(purchase)
        
        _remove_trailing_empty_paragraphs(doc)
        
        # Configure styles just in case for the appended documents
        _configure_document(doc)
        
        doc.add_page_break()
        # 2. ใบสั่งซื้อ (po)
        _add_purchase_order(doc, purchase)
        
        doc.add_page_break()
        # 3. แบบกำหนด Spec (spec)
        _add_spec(doc, purchase)
        
        doc.add_page_break()
        # 7. แบบแสดงความบริสุทธิ์ใจ (integrity)
        _add_integrity_form(doc, purchase)
        
        doc.add_page_break()
        # 6. ใบตรวจรับ (receipt)
        _add_acceptance_receipt(doc, purchase)
        
        doc.add_page_break()
        # 9. แบบฟอร์มกำหนด spec
        _add_spec_details(doc, purchase)
    else:
        raise ValueError("ไม่พบรูปแบบ Word")

    # Export every Word file without editing restrictions or locked objects.
    _make_word_fully_editable(doc)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output



# ---------------- Excel export: every Word table ----------------
def _build_document_for_excel_tables(purchase):
    """
    สร้าง document แบบเดียวกับ Word 'เอกสารทั้งหมด'
    แล้วนำทุก Word table ไป export เป็น Excel
    """
    with _DisableDigitFontSplit():
        doc = _build_exact_procurement_template(purchase)

    # เพิ่มตารางหน้าสุดท้ายด้วย เพื่อให้ Excel มีตารางนี้ด้วย
    _add_final_spec_excel_page(doc, purchase)

    return doc


def _excel_safe_sheet_name(name, used_names):
    invalid = '[]:*?/\\'
    cleaned = str(name or "Table")

    for ch in invalid:
        cleaned = cleaned.replace(ch, "_")

    cleaned = cleaned.strip() or "Table"
    cleaned = cleaned[:31]

    base = cleaned
    counter = 2

    while cleaned in used_names:
        suffix = f"_{counter}"
        cleaned = (base[:31 - len(suffix)] + suffix)
        counter += 1

    used_names.add(cleaned)
    return cleaned


def _table_title_from_word_table(table, table_no):
    """
    ตั้งชื่อ Sheet จากข้อความในตาราง ถ้าหาไม่ได้ใช้ Table_XX
    """
    texts = []

    for row in table.rows[:2]:
        for cell in row.cells:
            value = " ".join((cell.text or "").split())
            if value and value not in texts:
                texts.append(value)

    joined = " ".join(texts)

    keywords = [
        ("รายการพิจารณา", "ผลการพิจารณา"),
        ("ยอดที่ได้รับจัดสรร", "งบประมาณ"),
        ("ราคาต่อหน่วย", "ใบสั่งซื้อ"),
        ("หน่วยละ", "กำหนด_Spec"),
        ("คุณลักษณะเฉพาะ", "รายละเอียด_Spec"),
        ("ผลการตรวจรับ", "ตรวจรับ"),
    ]

    for keyword, title in keywords:
        if keyword in joined:
            return f"{table_no:02d}_{title}"

    if texts:
        first = texts[0][:18]
        return f"{table_no:02d}_{first}"

    return f"Table_{table_no:02d}"


def _write_word_table_to_worksheet(ws, table):
    """
    นำ Word table 1 ตารางลง Excel 1 Sheet
    พร้อมกรอบ เส้นตาราง การจัดกึ่งกลาง และ wrap text
    """
    thin = Side(style="thin", color="000000")
    full_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    header_fill = PatternFill(fill_type="solid", fgColor="EDEDED")
    header_font = Font(name="TH Sarabun New", size=14, bold=True)
    normal_font = Font(name="TH Sarabun New", size=13)

    # Write values.
    for row_index, row in enumerate(table.rows, start=1):
        for col_index, cell in enumerate(row.cells, start=1):
            value = (cell.text or "").strip()

            excel_cell = ws.cell(
                row=row_index,
                column=col_index,
                value=value,
            )

            excel_cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
                wrap_text=True,
            )
            excel_cell.border = full_border
            excel_cell.font = normal_font

            if row_index == 1:
                excel_cell.fill = header_fill
                excel_cell.font = header_font

    # Try to preserve horizontal merged cells from Word.
    for row_index, row in enumerate(table.rows, start=1):
        tc_ids = [id(cell._tc) for cell in row.cells]
        start_col = 0

        while start_col < len(tc_ids):
            end_col = start_col

            while (
                end_col + 1 < len(tc_ids)
                and tc_ids[end_col + 1] == tc_ids[start_col]
            ):
                end_col += 1

            if end_col > start_col:
                try:
                    ws.merge_cells(
                        start_row=row_index,
                        start_column=start_col + 1,
                        end_row=row_index,
                        end_column=end_col + 1,
                    )
                except Exception:
                    pass

            start_col = end_col + 1

    # Widths based on actual contents, but cap so sheets remain readable.
    max_columns = max((len(row.cells) for row in table.rows), default=1)

    for col_index in range(1, max_columns + 1):
        max_len = 0

        for row_index in range(1, len(table.rows) + 1):
            value = ws.cell(row=row_index, column=col_index).value
            if value is None:
                continue

            lines = str(value).splitlines() or [""]
            max_len = max(max_len, max(len(line) for line in lines))

        width = min(max(max_len + 3, 12), 45)
        ws.column_dimensions[get_column_letter(col_index)].width = width

    for row_index in range(1, len(table.rows) + 1):
        ws.row_dimensions[row_index].height = 26

    ws.freeze_panes = "A2"

    # Print setup.
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_options.horizontalCentered = True


def _build_excel_from_word_tables(purchase, only_table_no=None):
    """
    only_table_no=None -> ทุกตาราง แยกเป็นหลาย Sheet
    only_table_no=N    -> เฉพาะตารางที่ N
    """
    doc = _build_document_for_excel_tables(purchase)
    tables = list(doc.tables)

    if only_table_no is not None:
        if only_table_no < 1 or only_table_no > len(tables):
            raise ValueError("ไม่พบหมายเลขตาราง")
        selected = [(only_table_no, tables[only_table_no - 1])]
    else:
        selected = list(enumerate(tables, start=1))

    workbook = ExcelWorkbook()

    # ลบ Sheet เปล่าเริ่มต้น
    default_sheet = workbook.active
    workbook.remove(default_sheet)

    used_names = set()

    for table_no, table in selected:
        suggested = _table_title_from_word_table(table, table_no)
        sheet_name = _excel_safe_sheet_name(suggested, used_names)
        ws = workbook.create_sheet(title=sheet_name)

        _write_word_table_to_worksheet(ws, table)

    output = BytesIO()
    workbook.save(output)
    output.seek(0)

    return output, len(tables)


@main_bp.route("/purchases/<int:purchase_id>/excel/all")
def export_excel_all_tables(purchase_id):
    """
    ดาวน์โหลด Excel 1 ไฟล์:
    ทุกตารางในเอกสารแยกเป็นคนละ Sheet
    """
    purchase = Purchase.query.get_or_404(purchase_id)

    output, table_count = _build_excel_from_word_tables(
        purchase,
        only_table_no=None,
    )

    filename = safe_filename(
        f"{purchase.po_number}_Excel_ทุกตาราง_{table_count}_ตาราง.xlsx"
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        max_age=0,
    )


@main_bp.route("/purchases/<int:purchase_id>/excel/table/<int:table_no>")
def export_excel_single_table(purchase_id, table_no):
    """
    ดาวน์โหลด Excel แยกทีละตาราง
    เช่น /purchases/2/excel/table/1
    """
    purchase = Purchase.query.get_or_404(purchase_id)

    try:
        output, table_count = _build_excel_from_word_tables(
            purchase,
            only_table_no=table_no,
        )
    except ValueError:
        return "ไม่พบตาราง", 404

    filename = safe_filename(
        f"{purchase.po_number}_ตาราง_{table_no:02d}_จาก_{table_count}.xlsx"
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        max_age=0,
    )


@main_bp.route("/purchases/<int:purchase_id>/excel/tables")
def excel_table_index(purchase_id):
    """
    JSON สำหรับดูว่ามีกี่ตาราง และ URL ดาวน์โหลดแต่ละตาราง
    """
    purchase = Purchase.query.get_or_404(purchase_id)
    doc = _build_document_for_excel_tables(purchase)

    result = []

    for table_no, table in enumerate(doc.tables, start=1):
        result.append({
            "table_no": table_no,
            "name": _table_title_from_word_table(table, table_no),
            "download_url": url_for(
                "main.export_excel_single_table",
                purchase_id=purchase.id,
                table_no=table_no,
            ),
        })

    return jsonify({
        "purchase_id": purchase.id,
        "po_number": purchase.po_number,
        "table_count": len(result),
        "download_all_url": url_for(
            "main.export_excel_all_tables",
            purchase_id=purchase.id,
        ),
        "tables": result,
    })




@main_bp.route("/purchases/<int:purchase_id>/word/<form_type>")
def export_word(purchase_id, form_type):
    purchase = Purchase.query.get_or_404(purchase_id)
    allowed = {
        "po_single", "po_multiple", "spec",
        "receipt_single", "receipt_two", "receipt_multiple",
        "integrity", "procurement_pack", "all",
        # backwards-compatible names
        "po", "receipt",
    }
    if form_type not in allowed:
        return "ไม่พบแบบฟอร์ม", 404

    output = _build_word(purchase, form_type)
    suffixes = {
        "po_single": "ใบสั่งซื้อ_1รายการ",
        "po_multiple": "ใบสั่งซื้อ_หลายรายการ",
        "po": "ใบสั่งซื้อ",
        "spec": "แบบกำหนดสเปค",
        "receipt_single": "ใบตรวจรับ_1รายการ",
        "receipt_two": "ใบตรวจรับ_2รายการ",
        "receipt_multiple": "ใบตรวจรับ_มากกว่า2รายการ",
        "receipt": "ใบตรวจรับ",
        "integrity": "แบบแสดงความบริสุทธิ์ใจ",
        "procurement_pack": "ชุดรายงานจัดซื้อ",
        "all": "เอกสารทั้งหมด",
    }
    filename = safe_filename(f"{purchase.po_number}_{suffixes[form_type]}.docx")
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        max_age=0,
    )
